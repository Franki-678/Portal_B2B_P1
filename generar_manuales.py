#!/usr/bin/env python3
"""
generar_manuales.py — Genera los 6 documentos oficiales de Portal B2B · RC Repuestos
y los convierte a PDF.

Uso:
    python generar_manuales.py

Salida:
    Entrega_Final/
        01_Resumen_Ejecutivo.docx + .pdf
        02_Guia_Admin.docx        + .pdf
        03_Guia_Vendedores.docx   + .pdf
        04_Guia_Talleres.docx     + .pdf
        05_Acuerdo_SLA.docx       + .pdf
        06_Handover_Interno.docx  + .pdf
"""

import os
import sys
import shutil
from pathlib import Path

# Forzar UTF-8 en la consola de Windows
if sys.platform == "win32":
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─────────────────────────────────────────────────────────────────────────────
# CONSTANTES GLOBALES
# ─────────────────────────────────────────────────────────────────────────────
EMPRESA        = "RC Repuestos"
ADMIN_NOMBRE   = "Pagliero Juan Ludovico"
CONTACTO_SOPORTE = (
    "Luis Giannasi — ManIAco\n"
    "Teléfono / WhatsApp: 3516009131\n"
    "Mail: luis.giannasi@maniaco.online"
)
CONTACTO_INLINE = "Luis Giannasi (ManIAco) · 3516009131 · luis.giannasi@maniaco.online"
APP_URL         = "https://portal-b2b.vercel.app"
FECHA_HOY       = datetime.date.today().strftime("%-d de %B de %Y") if sys.platform != "win32" else datetime.date.today().strftime("%d/%m/%Y")
OUTPUT_DIR      = Path(__file__).parent / "Entrega_Final"


# ─────────────────────────────────────────────────────────────────────────────
# HELPERS DE FORMATO
# ─────────────────────────────────────────────────────────────────────────────

def new_doc(title: str, subtitle: str = "") -> Document:
    """Crea un documento con portada estándar."""
    doc = Document()
    # Márgenes
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2.5)
    # Título principal
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.runs[0]
    run.font.size  = Pt(20)
    run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
    # Subtítulo
    if subtitle:
        p = doc.add_paragraph(subtitle)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size  = Pt(12)
        p.runs[0].font.color.rgb = RGBColor(0x6b, 0x72, 0x80)
    # Empresa y fecha
    meta = doc.add_paragraph(f"{EMPRESA}  ·  Generado: {FECHA_HOY}")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].font.size  = Pt(10)
    meta.runs[0].font.color.rgb = RGBColor(0x9c, 0xa3, 0xaf)
    doc.add_paragraph("")
    return doc


def h1(doc: Document, text: str):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
    return p


def h2(doc: Document, text: str):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
    return p


def h3(doc: Document, text: str):
    return doc.add_heading(text, level=3)


def body(doc: Document, text: str):
    return doc.add_paragraph(text)


def bullet(doc: Document, text: str, level: int = 0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(level * 0.25)
    return p


def numbered(doc: Document, text: str):
    return doc.add_paragraph(text, style="List Number")


def bold_label(doc: Document, label: str, value: str):
    """Línea con etiqueta en negrita seguida de valor normal."""
    p = doc.add_paragraph()
    run_l = p.add_run(f"{label}: ")
    run_l.bold = True
    p.add_run(value)
    return p


def divider(doc: Document):
    doc.add_paragraph("─" * 60)


def contacto_footer(doc: Document):
    doc.add_paragraph("")
    divider(doc)
    h3(doc, "Contacto de soporte técnico")
    body(doc, CONTACTO_SOPORTE)


def add_table(doc: Document, headers: list, rows: list):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
    # Data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
    doc.add_paragraph("")
    return table


# ─────────────────────────────────────────────────────────────────────────────
# DOC 01 — RESUMEN EJECUTIVO
# ─────────────────────────────────────────────────────────────────────────────

def crear_resumen_ejecutivo() -> Path:
    doc = new_doc(
        "Portal B2B · RC Repuestos",
        "Resumen Ejecutivo del Sistema"
    )

    h1(doc, "¿Qué es Portal B2B?")
    body(doc,
        "Portal B2B es una plataforma web privada diseñada exclusivamente para RC Repuestos. "
        "Permite gestionar de punta a punta el circuito de pedidos de repuestos entre los talleres "
        "clientes y el equipo de vendedores de RC Repuestos, reemplazando el seguimiento manual "
        "por WhatsApp o teléfono."
    )
    body(doc,
        "Todo queda registrado: quién pidió qué, quién cotizó, cuánto se aprobó, cuándo se entregó "
        "y si hubo algún reclamo."
    )

    h1(doc, "¿Para quién es?")
    add_table(doc,
        ["Perfil", "Qué puede hacer"],
        [
            ["Administrador (Juan)", "Ver todo, crear vendedores, resolver conflictos, consultar métricas vía Telegram"],
            ["Vendedor", "Tomar pedidos, enviar cotizaciones, registrar pagos y entregas"],
            ["Taller", "Crear pedidos, aprobar cotizaciones, iniciar reclamos, subir fotos"],
        ]
    )

    h1(doc, "Ciclo de vida de un pedido")
    body(doc,
        "Cada pedido pasa por estados bien definidos. A continuación se describe el flujo completo:"
    )

    add_table(doc,
        ["Estado", "Significado", "Quién actúa"],
        [
            ["Pendiente",          "El taller creó el pedido, esperando que un vendedor lo tome",        "Vendedor"],
            ["En revisión",        "Un vendedor tomó el pedido y está preparando cotización",            "Vendedor"],
            ["Cotizado",           "El vendedor envió la cotización al taller",                          "Taller"],
            ["Aprobado",           "El taller aprobó toda la cotización",                                "Vendedor"],
            ["Aprobado parcial",   "El taller aprobó solo algunos ítems",                               "Vendedor"],
            ["Pagado · Por entregar", "El vendedor registró el cobro, falta entregar la mercadería",    "Vendedor"],
            ["Rechazado",          "El taller rechazó la cotización",                                   "Vendedor / Admin"],
            ["Cerrado · Pagado",   "Pedido entregado y cobrado. Ciclo completo",                        "—"],
            ["En conflicto",       "El taller inició un reclamo después de la entrega",                 "Admin"],
            ["Cancelado",          "El pedido fue cancelado antes de cerrar",                           "Admin"],
        ]
    )

    h2(doc, "Flujo principal (camino feliz)")
    numbered(doc, "Taller crea el pedido → estado: Pendiente")
    numbered(doc, "Vendedor lo toma → estado: En revisión")
    numbered(doc, "Vendedor cotiza → estado: Cotizado")
    numbered(doc, "Taller aprueba → estado: Aprobado o Aprobado parcial")
    numbered(doc, "Vendedor confirma cobro → estado: Pagado · Por entregar")
    numbered(doc, "Vendedor confirma entrega → estado: Cerrado · Pagado  ✓")

    h2(doc, "Caminos alternativos")
    bullet(doc, "Si el taller rechaza la cotización → estado: Rechazado. El pedido queda cerrado. "
                "Solo si el vendedor y el taller acordaron por WhatsApp que se va a volver a cotizar, "
                "el vendedor puede usar 'Re-cotizar (Preacordado)' para abrir el pedido nuevamente.")
    bullet(doc, "Si el taller inicia un reclamo después de recibir la mercadería → estado: En conflicto. "
                "Juan resuelve el conflicto desde el panel de administración.")

    h1(doc, "Cómo los conflictos afectan los números")
    body(doc,
        "Cuando Juan resuelve un conflicto puede indicar que se hizo un descuento o una devolución parcial. "
        "En ese caso, el sistema registra el monto ajustado (el importe del descuento o devolución) "
        "vinculado al pedido. Ese ajuste se muestra en el detalle del pedido y en el panel de métricas, "
        "de modo que los KPIs (facturación total, ticket promedio) reflejen el valor neto real cobrado, "
        "no el valor bruto de la cotización original."
    )
    body(doc,
        "Ejemplo: si un pedido fue cotizado por $100.000 pero se resolvió con un descuento de $15.000, "
        "el sistema registra el ajuste y el importe neto efectivo es $85.000."
    )

    h1(doc, "¿Qué hace el sistema por RC Repuestos?")
    bullet(doc, "Elimina el seguimiento manual por WhatsApp o planilla")
    bullet(doc, "Todos los pedidos, cotizaciones y pagos quedan registrados con fecha y hora")
    bullet(doc, "Juan recibe métricas del negocio en tiempo real por Telegram (sin abrir el browser)")
    bullet(doc, "Los talleres pueden iniciar reclamos con fotos adjuntas")
    bullet(doc, "Los vendedores ven el método de pago elegido por el taller directamente en la plataforma")
    bullet(doc, "Notificaciones automáticas al grupo de Telegram de ventas en cada evento importante")

    h1(doc, "Infraestructura y costos")
    body(doc,
        "El sistema corre sobre capas gratuitas de proveedores en la nube. No hay costos fijos mensuales "
        "actualmente."
    )
    add_table(doc,
        ["Componente", "Proveedor", "Plan"],
        [
            ["Sitio web y API",   "Vercel",   "Hobby (gratuito)"],
            ["Base de datos",     "Supabase", "Free Tier (gratuito)"],
            ["Bot de Telegram",   "Telegram", "Gratuito"],
            ["Notificaciones",    "Telegram", "Gratuito — sin mail (Resend eliminado)"],
        ]
    )
    body(doc,
        "Existe un período de prueba de 1 mes para detectar y corregir errores sin costo. "
        "Al finalizar ese período se define comercialmente si se migra a planes pagos según el uso."
    )

    contacto_footer(doc)
    path = OUTPUT_DIR / "01_Resumen_Ejecutivo.docx"
    doc.save(path)
    print(f"  ✓ {path.name}")
    return path


# ─────────────────────────────────────────────────────────────────────────────
# DOC 02 — GUÍA ADMINISTRADOR
# ─────────────────────────────────────────────────────────────────────────────

def crear_guia_admin() -> Path:
    doc = new_doc(
        "Guía del Administrador",
        f"Portal B2B · RC Repuestos — Uso exclusivo de {ADMIN_NOMBRE}"
    )

    h1(doc, "Tu rol como administrador")
    body(doc,
        "Como administrador tenés acceso completo al sistema. Sos la única persona que puede "
        "crear cuentas de vendedor, resolver conflictos entre talleres y vendedores, y consultar "
        "las métricas del negocio."
    )

    h1(doc, "Acceso al sistema")
    bold_label(doc, "URL", APP_URL)
    bold_label(doc, "Usuario", "Tu correo electrónico registrado")
    bold_label(doc, "Contraseña", "La que configuraste al activar tu cuenta")
    body(doc,
        "No existe recuperación por mail automática (Resend fue eliminado). "
        "Si perdés el acceso, contactá al soporte técnico."
    )

    h1(doc, "Cómo crear una cuenta de vendedor")
    body(doc,
        "Los vendedores NO pueden registrarse solos. Vos creás su cuenta manualmente. "
        "No se envía ningún mail de aviso: tenés que comunicarle las credenciales al vendedor "
        "directamente (por WhatsApp, teléfono, o como prefieras)."
    )
    numbered(doc, "Ingresá al panel de administración → sección 'Vendedores'.")
    numbered(doc, "Hacé clic en 'Nuevo vendedor'.")
    numbered(doc, "Completá: nombre completo, correo electrónico y contraseña temporal.")
    numbered(doc, "Opcionalmente: ingresá el usuario de Telegram del vendedor (sin '@') para que el bot "
                  "lo etiquete en los mensajes del grupo cuando hay una acción que requiere su atención.")
    numbered(doc, "Guardá. La cuenta queda activa de inmediato.")
    numbered(doc, "Avisale al vendedor por WhatsApp o teléfono: su mail y contraseña temporal. "
                  "Pedile que cambie la contraseña en su primer ingreso.")
    body(doc,
        "Importante: si el vendedor no tiene usuario de Telegram cargado, el bot igual funciona "
        "pero en lugar de etiquetarlo muestra su nombre en negrita. "
        "El bot NUNCA etiqueta ('pingea') a nadie que no tenga el username cargado."
    )

    h1(doc, "Gestión del panel de pedidos")
    body(doc,
        "Desde tu panel podés ver TODOS los pedidos, sin importar a qué vendedor estén asignados. "
        "Los pedidos aparecen ordenados del más nuevo al más antiguo."
    )
    bullet(doc, "Podés filtrar por estado, taller o vendedor")
    bullet(doc, "Podés entrar al detalle de cualquier pedido y ver el historial completo de eventos")
    bullet(doc, "Podés ver el método de pago que eligió el taller (transferencia o efectivo)")
    bullet(doc, "Podés ver si hay un ajuste de monto por conflicto resuelto")

    h1(doc, "Resolución de conflictos")
    body(doc,
        "Cuando un taller inicia un reclamo, el pedido pasa a estado 'En conflicto'. "
        "Ese pedido aparece resaltado en rojo en tu panel y el bot te avisa por Telegram."
    )
    numbered(doc, "Ingresá al pedido en conflicto.")
    numbered(doc, "Leé el comentario del taller (y las fotos adjuntas si las hay).")
    numbered(doc, "Hablar con el vendedor y el taller para acordar una solución.")
    numbered(doc, "En el panel, hacé clic en 'Resolver conflicto'.")
    numbered(doc, "Elegí el resultado:")
    bullet(doc, "Resuelto sin ajuste: todo OK, sin descuentos", level=1)
    bullet(doc, "Descuento: ingresá el monto o porcentaje de descuento", level=1)
    bullet(doc, "Devolución parcial: ingresá el monto a devolver", level=1)
    bullet(doc, "Reposición: se reenvía mercadería sin costo", level=1)
    numbered(doc, "Escribí una nota explicando lo acordado.")
    numbered(doc, "Confirmá. El pedido vuelve a 'Cerrado · Pagado' y el ajuste queda registrado.")

    h1(doc, "Comandos de Telegram (solo para vos)")
    body(doc,
        "El bot responde a comandos ÚNICAMENTE si los enviás desde el chat que tenés configurado "
        "como administrador. Cualquier otro usuario que intente usarlos no recibe respuesta."
    )
    add_table(doc,
        ["Comando", "Qué devuelve"],
        [
            ["/hoy",       "Facturado hoy vs. ayer, pedidos entregados, nuevos pedidos, pendientes y conflictos activos"],
            ["/vendedores","Ranking mensual de vendedores por facturación, comparativa con el mes anterior y resumen de la semana actual"],
            ["/alertas",   "Pedidos sin vendedor asignado, conflictos activos y pedidos trabados hace más de 48 horas, todos con enlace directo al pedido"],
        ]
    )
    body(doc,
        "Los números de pedido en los mensajes del bot son enlaces clickeables que te llevan "
        "directamente al detalle del pedido en la plataforma."
    )

    h1(doc, "Altas de taller")
    body(doc,
        "Los talleres pueden registrarse solos en la plataforma usando el formulario de registro. "
        "No necesitan pedirle cuenta a ningún vendedor. Una vez registrados, pueden crear pedidos "
        "de inmediato."
    )

    h1(doc, "Glosario rápido")
    add_table(doc,
        ["Término", "Significado en el sistema"],
        [
            ["Etiquetar",   "El bot menciona al vendedor en el mensaje de Telegram (ej: @juan_p)"],
            ["Deep link",   "Enlace que te lleva directo al pedido dentro de la plataforma"],
            ["KPI",         "Número clave del negocio: facturación, entregas, ticket promedio"],
            ["Ajuste",      "Descuento o devolución aplicada al cerrar un conflicto"],
            ["Preacordado", "Re-cotización acordada manualmente (fuera del sistema) con el taller"],
        ]
    )

    contacto_footer(doc)
    path = OUTPUT_DIR / "02_Guia_Admin.docx"
    doc.save(path)
    print(f"  ✓ {path.name}")
    return path


# ─────────────────────────────────────────────────────────────────────────────
# DOC 03 — GUÍA VENDEDORES
# ─────────────────────────────────────────────────────────────────────────────

def crear_guia_vendedores() -> Path:
    doc = new_doc(
        "Guía del Vendedor",
        "Portal B2B · RC Repuestos"
    )

    h1(doc, "Cómo acceder")
    bold_label(doc, "URL", APP_URL)
    body(doc,
        "Tu cuenta fue creada por Juan (el administrador). Él te dio tu mail y contraseña temporal. "
        "Al ingresar por primera vez, cambiá la contraseña desde tu perfil."
    )
    body(doc,
        "No existe recuperación automática de contraseña por mail. "
        "Si la olvidás, avisale a Juan para que la resetee."
    )

    h1(doc, "Tu panel de pedidos")
    body(doc,
        "Al ingresar ves la lista de todos los pedidos disponibles y los que tenés asignados. "
        "Los pedidos aparecen ordenados del más nuevo al más antiguo, así siempre "
        "ves primero los que acabaron de entrar."
    )
    body(doc,
        "Podés filtrar por estado (pendiente, en revisión, cotizado, etc.) para encontrar "
        "rápidamente lo que necesitás."
    )

    h1(doc, "Tomar un pedido")
    numbered(doc, "Buscá un pedido en estado 'Pendiente'.")
    numbered(doc, "Entrá al detalle y hacé clic en 'Tomar pedido'.")
    numbered(doc, "El pedido pasa a 'En revisión' y queda asignado a vos.")
    numbered(doc, "El grupo de Telegram recibe una notificación automática.")
    body(doc,
        "Una vez que tomás un pedido, otros vendedores no pueden tomarlo. "
        "Si no lo podés atender, podés liberarlo para que otro lo tome."
    )

    h1(doc, "Enviar una cotización")
    numbered(doc, "Entrá al pedido que tenés en 'En revisión'.")
    numbered(doc, "Completá precio y cantidad disponible para cada ítem solicitado.")
    numbered(doc, "Hacé clic en 'Enviar cotización'.")
    numbered(doc, "El pedido pasa a 'Cotizado' y el taller puede verla.")
    body(doc,
        "Tip: si no tenés un ítem, podés cotizarlo con precio 0 o dejar una nota en el comentario."
    )

    h1(doc, "Qué hacer cuando el taller aprueba")
    body(doc,
        "Cuando el taller aprueba (total o parcialmente), el pedido pasa a 'Aprobado' o "
        "'Aprobado parcial'. El bot de Telegram te etiqueta en el grupo para avisarte."
    )
    numbered(doc, "Coordiná el cobro con el taller.")
    numbered(doc, "Una vez cobrado, marcá el pedido como pagado en la plataforma.")
    numbered(doc, "Llevá la mercadería.")
    numbered(doc, "Confirmá la entrega en la plataforma → el pedido cierra como 'Cerrado · Pagado'.")
    body(doc,
        "Importante: el taller puede haber elegido pagar por transferencia o en efectivo. "
        "Podés ver el método de pago elegido directamente en el detalle del pedido, "
        "dentro de la plataforma, sin necesidad de consultarle."
    )

    h1(doc, "Qué hacer cuando el taller rechaza la cotización")
    body(doc,
        "Si el taller rechaza tu cotización, el pedido pasa a estado 'Rechazado'. "
        "Ese pedido queda cerrado. No podés volver a cotizar automáticamente."
    )
    body(doc,
        "Si hablaste con el taller por WhatsApp y acordaron que vas a volver a cotizar con "
        "otros precios o condiciones, podés usar el botón 'Re-cotizar (Preacordado)':"
    )
    numbered(doc, "Entrá al pedido en estado 'Rechazado'.")
    numbered(doc, "Hacé clic en 'Re-cotizar (Preacordado)'.")
    numbered(doc, "Confirmá. La cotización anterior se borra y el pedido vuelve a 'En revisión'.")
    numbered(doc, "Podés enviar una nueva cotización desde cero.")
    body(doc,
        "Este botón existe para cuando YA acordaste con el taller que vas a re-cotizar. "
        "No lo uses si el taller no te lo pidió."
    )

    h1(doc, "Notificaciones de Telegram")
    body(doc,
        "El sistema envía mensajes automáticos al grupo de ventas de Telegram cada vez que:"
    )
    bullet(doc, "Se crea un pedido nuevo")
    bullet(doc, "Tomás un pedido")
    bullet(doc, "Enviás una cotización")
    bullet(doc, "Un taller aprueba o rechaza tu cotización (el bot te etiqueta)")
    bullet(doc, "Registrás un pago o una entrega")
    bullet(doc, "Un taller inicia un reclamo (el bot te etiqueta)")
    body(doc,
        "No hay notificaciones por mail. Todo pasa por Telegram."
    )

    h1(doc, "Preguntas frecuentes")
    h3(doc, "¿Puedo ver pedidos de otros vendedores?")
    body(doc, "Podés ver los pedidos disponibles (sin asignar), pero no el detalle de los pedidos "
              "que ya tomó otro vendedor.")
    h3(doc, "¿Qué pasa si entrego sin registrarlo en el sistema?")
    body(doc, "El pedido queda en 'Pagado · Por entregar' hasta que lo confirmes. "
              "Siempre registrá la entrega en la plataforma para que los números cierren.")
    h3(doc, "¿Puedo cancelar un pedido?")
    body(doc, "Solo Juan puede cancelar pedidos.")

    contacto_footer(doc)
    path = OUTPUT_DIR / "03_Guia_Vendedores.docx"
    doc.save(path)
    print(f"  ✓ {path.name}")
    return path


# ─────────────────────────────────────────────────────────────────────────────
# DOC 04 — GUÍA TALLERES
# ─────────────────────────────────────────────────────────────────────────────

def crear_guia_talleres() -> Path:
    doc = new_doc(
        "Guía del Taller",
        "Portal B2B · RC Repuestos"
    )

    h1(doc, "¿Qué es Portal B2B?")
    body(doc,
        "Es la plataforma en línea de RC Repuestos para hacer pedidos de repuestos. "
        "Desde acá podés pedir, ver las cotizaciones, aprobarlas y hacer reclamos, "
        "todo en un solo lugar y sin necesidad de llamar o mandar mensajes."
    )

    h1(doc, "Cómo registrarse")
    body(doc,
        "Podés crear tu propia cuenta sin pedirle nada a nadie. El registro es libre:"
    )
    numbered(doc, f"Entrá a {APP_URL}")
    numbered(doc, "Hacé clic en 'Registrarse como taller'.")
    numbered(doc, "Completá: nombre del taller, número de taller, correo electrónico y contraseña.")
    numbered(doc, "Listo. Podés empezar a pedir de inmediato.")
    body(doc,
        "Si ya tenés cuenta y olvidaste la contraseña, contactá al soporte técnico para resetearla. "
        "No hay recuperación automática por mail."
    )

    h1(doc, "Cómo hacer un pedido")
    numbered(doc, "Ingresá a tu cuenta.")
    numbered(doc, "Hacé clic en 'Nuevo pedido'.")
    numbered(doc, "Completá los datos del vehículo: marca, modelo y año.")
    numbered(doc, "Elegí la calidad que necesitás: alta (original/premium), media (aftermarket) o económica.")
    numbered(doc, "Elegí cómo vas a pagar: transferencia bancaria o efectivo.")
    numbered(doc, "Listá los repuestos que necesitás, con descripción clara.")
    numbered(doc, "Opcionalmente: subí fotos o documentos de referencia.")
    numbered(doc, "Enviá el pedido. Queda en 'Pendiente' hasta que un vendedor lo tome.")
    body(doc,
        "Tip: cuanto más detallado sea el pedido (número de parte, medidas, marca preferida), "
        "más rápido y precisa va a ser la cotización."
    )

    h1(doc, "Cómo ver y aprobar una cotización")
    numbered(doc, "Cuando el vendedor envíe la cotización, el pedido aparece en estado 'Cotizado'.")
    numbered(doc, "Entrá al pedido y revisá los precios y cantidades.")
    numbered(doc, "Podés aprobar todo, aprobar solo los ítems que te interesan, o rechazar.")
    numbered(doc, "Si aprobás, escribí el método de pago si no lo hiciste antes.")
    body(doc,
        "Una vez que aprobás, el vendedor coordinará el cobro y la entrega con vos."
    )

    h1(doc, "Cómo iniciar un reclamo")
    body(doc,
        "Si recibiste la mercadería pero hay un problema (ítems faltantes, dañados, equivocados), "
        "podés iniciar un reclamo directamente desde el pedido entregado:"
    )
    numbered(doc, "Entrá al pedido en estado 'Cerrado · Pagado'.")
    numbered(doc, "Hacé clic en 'Iniciar reclamo'.")
    numbered(doc, "Describí el problema en el campo de texto.")
    numbered(doc, "Para subir fotos:")
    bullet(doc, "Hacé clic en el botón de adjuntar (ícono de clip o cámara)", level=1)
    bullet(doc, "Seleccioná una o más imágenes desde tu dispositivo", level=1)
    bullet(doc, "Las fotos se suben automáticamente al guardar el reclamo", level=1)
    numbered(doc, "Confirmá. El pedido pasa a 'En conflicto' y Juan recibe una notificación por Telegram.")
    body(doc,
        "Juan revisará el reclamo y se comunicará con vos y con el vendedor para resolverlo."
    )

    h1(doc, "Método de pago")
    body(doc,
        "Cuando creás un pedido, elegís cómo vas a pagar: transferencia bancaria o efectivo. "
        "Esta información queda guardada en el pedido y el vendedor la puede ver directamente "
        "en la plataforma, sin necesidad de consultarte."
    )
    body(doc,
        "Si necesitás cambiar el método de pago después de crear el pedido, "
        "avisale al vendedor por WhatsApp para que lo coordinen."
    )

    h1(doc, "Notificaciones")
    body(doc,
        "El sistema NO envía notificaciones por mail. "
        "Para saber el estado de tu pedido, ingresá a la plataforma y revisá tu panel de pedidos. "
        "Todo está actualizado en tiempo real."
    )

    h1(doc, "Preguntas frecuentes")
    h3(doc, "¿Cuánto tarda en llegar una cotización?")
    body(doc, "Depende de la disponibilidad del vendedor. Normalmente entre algunas horas y un día hábil.")
    h3(doc, "¿Puedo pedir varios repuestos en un mismo pedido?")
    body(doc, "Sí, podés listar todos los repuestos que necesitás en un solo pedido.")
    h3(doc, "¿Qué pasa si no apruebo ni rechazo la cotización?")
    body(doc, "El pedido queda en 'Cotizado' hasta que tomés una decisión. No vence automáticamente.")
    h3(doc, "¿Puedo hacer un pedido desde el celular?")
    body(doc, "Sí. La plataforma está adaptada para móviles.")

    contacto_footer(doc)
    path = OUTPUT_DIR / "04_Guia_Talleres.docx"
    doc.save(path)
    print(f"  ✓ {path.name}")
    return path


# ─────────────────────────────────────────────────────────────────────────────
# DOC 05 — ACUERDO SLA
# ─────────────────────────────────────────────────────────────────────────────

def crear_acuerdo_sla() -> Path:
    doc = new_doc(
        "Acuerdo de Nivel de Servicio (SLA)",
        f"Portal B2B · RC Repuestos — Versión inicial · {FECHA_HOY}"
    )

    h1(doc, "1. Partes del acuerdo")
    add_table(doc,
        ["Rol", "Nombre / Empresa", "Contacto"],
        [
            ["Cliente",           f"{EMPRESA} — {ADMIN_NOMBRE}", "—"],
            ["Proveedor técnico", "Luis Giannasi — ManIAco",     "3516009131 · luis.giannasi@maniaco.online"],
        ]
    )

    h1(doc, "2. Alcance del servicio")
    body(doc,
        "Este acuerdo cubre el mantenimiento, soporte y disponibilidad de la plataforma "
        "Portal B2B desarrollada para RC Repuestos, que incluye:"
    )
    bullet(doc, "Sitio web y API alojados en Vercel")
    bullet(doc, "Base de datos alojada en Supabase")
    bullet(doc, "Bot de Telegram para notificaciones y comandos de gestión")
    bullet(doc, "Panel de administración, gestión de pedidos, cotizaciones y conflictos")

    h1(doc, "3. Infraestructura y costos actuales")
    body(doc,
        "El sistema opera sobre capas gratuitas de proveedores en la nube. "
        "No existen costos fijos mensuales en la configuración actual."
    )
    add_table(doc,
        ["Servicio", "Proveedor", "Plan", "Costo mensual"],
        [
            ["Hosting web y API", "Vercel",   "Hobby (Free)",    "$0"],
            ["Base de datos",     "Supabase", "Free Tier",       "$0"],
            ["Notificaciones",    "Telegram", "API gratuita",    "$0"],
            ["Mail transaccional","—",         "No aplica (eliminado)", "$0"],
        ]
    )
    body(doc,
        "Limitaciones del Free Tier a tener en cuenta:\n"
        "  · Vercel Hobby: proyectos no comerciales; 100 GB de ancho de banda/mes\n"
        "  · Supabase Free: 500 MB de base de datos; 2 GB de transferencia/mes; "
        "el proyecto se pausa si no tiene actividad por 7 días en el plan gratuito\n"
        "\n"
        "Si el volumen de uso supera estos límites, se migra a planes pagos. "
        "Esa decisión se toma al finalizar el período de prueba."
    )

    h1(doc, "4. Período de prueba inicial")
    body(doc,
        "Se establece un período de prueba de 1 (un) mes a partir de la puesta en producción. "
        "Durante ese mes:"
    )
    bullet(doc, "El proveedor técnico atiende bugs y errores de forma prioritaria sin costo adicional")
    bullet(doc, "Se detectan y corrigen inconsistencias de datos, errores de lógica y problemas de UX")
    bullet(doc, "No se cobran horas de soporte por correcciones derivadas del desarrollo inicial")
    body(doc,
        "Al finalizar el mes de prueba, las partes definen comercialmente las condiciones de "
        "mantenimiento y soporte continuo: tarifa mensual, horas incluidas, SLA de respuesta, etc."
    )

    h1(doc, "5. Disponibilidad")
    body(doc,
        "Durante el período de prueba, el objetivo de disponibilidad es del 95% mensual "
        "(aproximadamente 36 horas de caída toleradas por mes). Este objetivo está sujeto a las "
        "limitaciones de los proveedores gratuitos (Vercel y Supabase)."
    )
    body(doc,
        "Las interrupciones causadas por mantenimiento planificado, caídas de los proveedores "
        "de infraestructura o fuerza mayor no cuentan como tiempo de inactividad del SLA."
    )

    h1(doc, "6. Comandos de Telegram para el administrador")
    body(doc,
        "Juan (administrador) tiene acceso a tres comandos de gestión en el bot de Telegram. "
        "Estos comandos responden ÚNICAMENTE al chat de Juan; cualquier otro usuario recibe silencio."
    )
    add_table(doc,
        ["Comando", "Descripción"],
        [
            ["/hoy",
             "Resumen del día actual: facturación, pedidos entregados y nuevos, comparados con el día anterior. "
             "También muestra pendientes totales y conflictos activos."],
            ["/vendedores",
             "Ranking mensual de vendedores por facturación y entregas. Incluye comparativa con el mes anterior "
             "(crecimiento porcentual) y resumen de la semana en curso."],
            ["/alertas",
             "Lista de pedidos sin vendedor asignado, pedidos en conflicto activo (con enlace directo a cada uno) "
             "y pedidos trabados hace más de 48 horas sin movimiento."],
        ]
    )

    h1(doc, "7. Soporte y tiempos de respuesta")
    body(doc,
        "Durante el período de prueba (primer mes):"
    )
    add_table(doc,
        ["Tipo de incidencia", "Tiempo de respuesta objetivo"],
        [
            ["Error crítico (sistema caído, pérdida de datos)", "4 horas hábiles"],
            ["Error funcional (feature no funciona)",           "24 horas hábiles"],
            ["Consulta o mejora menor",                         "48 a 72 horas hábiles"],
        ]
    )
    body(doc,
        "Canal de contacto: WhatsApp al 3516009131 (Luis Giannasi). "
        "Mail alternativo: luis.giannasi@maniaco.online"
    )

    h1(doc, "8. Lo que NO cubre este acuerdo")
    bullet(doc, "Capacitación de usuarios finales más allá de los manuales entregados")
    bullet(doc, "Nuevas funcionalidades no previstas en el desarrollo original")
    bullet(doc, "Soporte a usuarios individuales (talleres, vendedores) — eso lo gestiona Juan internamente")
    bullet(doc, "Problemas causados por modificaciones al código realizadas por terceros")
    bullet(doc, "Caídas de Vercel, Supabase o Telegram fuera del control del proveedor técnico")

    h1(doc, "9. Definición comercial post-prueba")
    body(doc,
        "Al cumplirse el mes de prueba, las partes se reúnen para definir:\n"
        "  · Continuidad del servicio y tarifa mensual de mantenimiento\n"
        "  · Migración a planes pagos de Vercel y/o Supabase si el uso lo requiere\n"
        "  · Alcance de futuras mejoras y su cotización\n"
        "\n"
        "Si no se llega a un acuerdo comercial, el código fuente queda en propiedad de RC Repuestos "
        "y puede ser mantenido por otro proveedor."
    )

    contacto_footer(doc)
    path = OUTPUT_DIR / "05_Acuerdo_SLA.docx"
    doc.save(path)
    print(f"  ✓ {path.name}")
    return path


# ─────────────────────────────────────────────────────────────────────────────
# DOC 06 — HANDOVER INTERNO (SOLO PARA JUAN)
# ─────────────────────────────────────────────────────────────────────────────

def crear_handover_interno() -> Path:
    doc = new_doc(
        "Guía de Handover — USO INTERNO",
        f"Portal B2B · RC Repuestos — CONFIDENCIAL — Solo para {ADMIN_NOMBRE}"
    )

    p = doc.add_paragraph(
        "⚠️  DOCUMENTO CONFIDENCIAL. No compartir. "
        "Contiene credenciales y datos de acceso críticos del sistema."
    )
    p.runs[0].bold = True
    p.runs[0].font.color.rgb = RGBColor(0xdc, 0x26, 0x26)

    # ── SECCIÓN 1: Variables de entorno ──────────────────────────────────────
    h1(doc, "1. Variables de entorno del sistema")
    body(doc,
        "Estas variables están cargadas en Vercel → Settings → Environment Variables. "
        "Son los 'secretos' que conectan la app con los servicios externos. "
        "Si alguna se pierde o cambia, el sistema deja de funcionar hasta que se corrija."
    )

    add_table(doc,
        ["Variable", "Para qué sirve", "Nota"],
        [
            ["NEXT_PUBLIC_SUPABASE_URL",
             "URL pública del proyecto Supabase",
             "Visible en Supabase → Project Settings → API"],
            ["NEXT_PUBLIC_SUPABASE_ANON_KEY",
             "Clave pública de Supabase (acceso con RLS)",
             "Visible en Supabase → Project Settings → API"],
            ["SUPABASE_SERVICE_ROLE_KEY",
             "Clave secreta de Supabase (acceso total, sin RLS). Solo en el servidor.",
             "Visible en Supabase → Project Settings → API → service_role. NUNCA exponerla al cliente."],
            ["SUPABASE_WEBHOOK_SECRET",
             "Secreto para validar que los webhooks que llegan al servidor vienen de Supabase y no de terceros",
             "Valor actual: super_secreto_juan_123"],
            ["TELEGRAM_BOT_TOKEN",
             "Token del bot de Telegram. Lo generó BotFather.",
             "Formato: 123456789:AAH... — No compartir. Si se filtra, regenerar en BotFather."],
            ["TELEGRAM_ADMIN_ID",
             "ID numérico del chat del administrador. El bot solo responde comandos a este ID.",
             "Valor: -1003901633687 (ID negativo = grupo/canal)"],
            ["TELEGRAM_GROUP_ID",
             "ID del grupo de ventas donde el bot envía notificaciones automáticas",
             "Configurable en Vercel. Distinto al ADMIN_ID."],
            ["NEXT_PUBLIC_APP_URL",
             "URL pública de la app. Usada para generar deep links en los mensajes del bot.",
             f"Valor: {APP_URL}"],
        ]
    )

    body(doc,
        "Todas estas variables también deben estar en el archivo .env.local para desarrollo local. "
        "Ese archivo NUNCA se sube al repositorio de GitHub (está en .gitignore)."
    )

    # ── SECCIÓN 2: Transferencia Vercel ──────────────────────────────────────
    h1(doc, "2. Cómo transferir el proyecto de Vercel")
    body(doc,
        "El proyecto en Vercel está asociado a la cuenta del desarrollador (ManIAco). "
        "Para pasarlo a tu control tenés dos opciones:"
    )

    h2(doc, "Opción A — Transferir a tu cuenta personal de Vercel (recomendada para empezar)")
    numbered(doc, "Creá una cuenta en vercel.com con tu mail (si no tenés una).")
    numbered(doc, "Avisale a Luis para que inicie la transferencia: él va a Vercel → el proyecto → Settings → "
                  "Transfer Project → ingresa tu mail de Vercel.")
    numbered(doc, "Vercel te manda un mail para aceptar la transferencia. Aceptá desde tu cuenta.")
    numbered(doc, "El proyecto queda en tu cuenta. Las variables de entorno se mantienen.")
    numbered(doc, "Si usás el plan Hobby (gratuito), no hay costo.")
    body(doc, "Nota: en el plan Hobby de Vercel el dominio queda como *.vercel.app. "
              "Si querés un dominio propio (ej: rcrepuestos.com.ar), necesitás comprarlo y configurarlo "
              "en Vercel → tu proyecto → Settings → Domains.")

    h2(doc, "Opción B — Crear una organización en Vercel")
    numbered(doc, "En vercel.com, creá una organización ('Team') con el nombre que quieras.")
    numbered(doc, "Desde la org, invitá a Luis para que transfiera el proyecto a la org.")
    numbered(doc, "Una vez dentro de la org, podés invitar a otras personas con acceso limitado.")
    body(doc, "Nota: las organizaciones en Vercel requieren plan Pro ($20/mes). "
              "Recomendado solo si vas a tener múltiples personas gestionando el proyecto.")

    # ── SECCIÓN 3: Transferencia Supabase ────────────────────────────────────
    h1(doc, "3. Cómo transferir el proyecto de Supabase")
    body(doc,
        "El proyecto de base de datos está en la cuenta de Supabase del desarrollador. "
        "Para pasarlo a tu control:"
    )

    h2(doc, "Paso 1 — Crear tu cuenta y organización en Supabase")
    numbered(doc, "Entrá a supabase.com y creá una cuenta con tu mail.")
    numbered(doc, "En el dashboard, hacé clic en 'New organization'.")
    numbered(doc, "Poné el nombre de la organización (ej: RC Repuestos).")
    numbered(doc, "El plan gratuito (Free) alcanza para el volumen actual.")

    h2(doc, "Paso 2 — Transferir el proyecto")
    numbered(doc, "Avisale a Luis el nombre de tu organización en Supabase.")
    numbered(doc, "Luis va a Supabase → el proyecto → Settings → General → Transfer project.")
    numbered(doc, "Selecciona tu organización como destino y confirma.")
    numbered(doc, "El proyecto aparece en tu organización con todos los datos intactos.")
    numbered(doc, "Verificá que las variables de entorno en Vercel sigan apuntando a las URLs correctas "
                  "(NEXT_PUBLIC_SUPABASE_URL y las claves API). Si el proyecto se transfirió sin cambiar "
                  "la URL, no necesitás cambiar nada.")

    h2(doc, "Paso 3 — Regenerar claves (recomendado por seguridad)")
    numbered(doc, "Una vez que el proyecto esté en tu organización, podés regenerar las claves API.")
    numbered(doc, "En Supabase → Settings → API → 'Reveal' → copiá las nuevas claves.")
    numbered(doc, "Actualizá NEXT_PUBLIC_SUPABASE_ANON_KEY y SUPABASE_SERVICE_ROLE_KEY en Vercel.")
    numbered(doc, "También actualizá SUPABASE_WEBHOOK_SECRET si querés cambiarlo.")

    # ── SECCIÓN 4: Transferencia del Bot de Telegram ─────────────────────────
    h1(doc, "4. Cómo transferir el bot de Telegram")
    body(doc,
        "El bot de Telegram fue creado en la cuenta de Telegram del desarrollador usando @BotFather. "
        "Telegram no tiene una función de 'transferencia' de bots, pero se puede hacer fácilmente:"
    )

    h2(doc, "Opción A — Cambiar el token (recomendada)")
    numbered(doc, "Abrí Telegram y buscá @BotFather.")
    numbered(doc, "Enviá: /mybots")
    numbered(doc, "Seleccioná el bot del Portal B2B (el nombre que eligió Luis).")
    numbered(doc, "Hacé clic en 'API Token' → 'Revoke current token'.")
    numbered(doc, "BotFather genera un nuevo token. Copialo.")
    numbered(doc, "Actualizá la variable TELEGRAM_BOT_TOKEN en Vercel con el nuevo token.")
    numbered(doc, "Re-registrá el webhook: enviá este mensaje a tu bot de Telegram "
                  "(o pedíselo a Luis para que lo haga):")
    doc.add_paragraph(
        "curl \"https://api.telegram.org/bot<NUEVO_TOKEN>/setWebhook"
        "?url=https://portal-b2b.vercel.app/api/webhooks/telegram/bot"
        "&secret_token=<últimos-20-caracteres-del-nuevo-token>\""
    ).runs[0].font.name = "Courier New"
    numbered(doc, "Listo. El bot sigue siendo el mismo pero el token es tuyo.")
    body(doc, "Importante: al revocar el token, el bot deja de funcionar hasta que registres "
              "el nuevo token en Vercel. Hacé los pasos 5, 6, 7 seguidos para minimizar el downtime.")

    h2(doc, "Opción B — Crear un bot nuevo")
    numbered(doc, "En @BotFather, enviá /newbot.")
    numbered(doc, "Elegí nombre y username para el nuevo bot.")
    numbered(doc, "Copiá el token y actualizalo en Vercel (TELEGRAM_BOT_TOKEN).")
    numbered(doc, "Re-registrá el webhook como en la Opción A.")
    numbered(doc, "Agregá el nuevo bot al grupo de Telegram y dale permisos de escritura.")
    body(doc, "Desventaja: los usuarios del grupo ven un 'bot nuevo'. "
              "Si el grupo ya está acostumbrado al bot anterior, puede generar confusión.")

    # ── SECCIÓN 5: Checklist de handover ─────────────────────────────────────
    h1(doc, "5. Checklist de handover completo")
    add_table(doc,
        ["Tarea", "Responsable", "¿Listo?"],
        [
            ["Crear cuenta en Vercel",                              "Juan",  "☐"],
            ["Transferir proyecto Vercel a cuenta/org de Juan",     "Luis",  "☐"],
            ["Crear cuenta en Supabase",                            "Juan",  "☐"],
            ["Crear organización en Supabase",                      "Juan",  "☐"],
            ["Transferir proyecto Supabase a org de Juan",          "Luis",  "☐"],
            ["Verificar variables de entorno en Vercel",            "Juan",  "☐"],
            ["Revocar token viejo del bot y generar nuevo",         "Luis",  "☐"],
            ["Actualizar TELEGRAM_BOT_TOKEN en Vercel",             "Juan",  "☐"],
            ["Re-registrar webhook del bot",                        "Juan",  "☐"],
            ["Probar /hoy, /vendedores, /alertas en Telegram",      "Juan",  "☐"],
            ["Confirmar que notificaciones del grupo funcionan",    "Juan",  "☐"],
            ["Hacer un pedido de prueba end-to-end",                "Juan",  "☐"],
            ["Acceso al repositorio GitHub (leer/fork si aplica)",  "Luis",  "☐"],
        ]
    )

    body(doc,
        "Ante cualquier duda durante el proceso de transferencia, contactar a Luis Giannasi "
        "al 3516009131 (WhatsApp)."
    )

    contacto_footer(doc)
    path = OUTPUT_DIR / "06_Handover_Interno.docx"
    doc.save(path)
    print(f"  ✓ {path.name}")
    return path


# ─────────────────────────────────────────────────────────────────────────────
# CONVERSIÓN A PDF
# ─────────────────────────────────────────────────────────────────────────────

def convertir_a_pdf(docx_paths: list) -> list:
    """Convierte .docx a PDF usando docx2pdf (requiere Microsoft Word en Windows)."""
    try:
        from docx2pdf import convert
        print("\n📄 Convirtiendo a PDF con docx2pdf...")
        pdfs = []
        for docx_path in docx_paths:
            pdf_path = docx_path.with_suffix(".pdf")
            try:
                convert(str(docx_path), str(pdf_path))
                if pdf_path.exists():
                    print(f"  ✓ {pdf_path.name}")
                    pdfs.append(pdf_path)
                else:
                    print(f"  ✗ No se generó: {pdf_path.name}")
            except Exception as e:
                print(f"  ✗ Error en {docx_path.name}: {e}")
        return pdfs
    except ImportError:
        print("  ✗ docx2pdf no disponible. Solo se generaron los archivos .docx")
        return []
    except Exception as e:
        print(f"  ✗ Error general en conversión PDF: {e}")
        print("     Instalá Microsoft Word o LibreOffice para habilitar la conversión.")
        return []


# ─────────────────────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────────────────────

def main():
    print()
    print("═" * 60)
    print("  PORTAL B2B · RC REPUESTOS — Generador de Manuales")
    print("═" * 60)
    print()

    # 1. Limpiar y crear directorio de salida
    if OUTPUT_DIR.exists():
        shutil.rmtree(OUTPUT_DIR)
        print(f"  ♻️  Carpeta anterior eliminada: {OUTPUT_DIR.name}/")
    OUTPUT_DIR.mkdir(parents=True)
    print(f"  📁 Carpeta creada: {OUTPUT_DIR.name}/")
    print()

    # 2. Generar documentos Word
    print("📝 Generando documentos Word...")
    docx_paths = []
    for fn in [
        crear_resumen_ejecutivo,
        crear_guia_admin,
        crear_guia_vendedores,
        crear_guia_talleres,
        crear_acuerdo_sla,
        crear_handover_interno,
    ]:
        try:
            path = fn()
            docx_paths.append(path)
        except Exception as e:
            print(f"  ✗ Error generando {fn.__name__}: {e}")

    # 3. Convertir a PDF
    pdfs = convertir_a_pdf(docx_paths)

    # 4. Resumen
    print()
    print("─" * 60)
    print(f"  📄 Word generados:  {len(docx_paths)}")
    print(f"  📑 PDF generados:   {len(pdfs)}")
    print(f"  📂 Destino:         {OUTPUT_DIR.resolve()}")
    print()

    # Listar archivos finales
    archivos = sorted(OUTPUT_DIR.iterdir())
    print("  Archivos en Entrega_Final/:")
    for f in archivos:
        size_kb = f.stat().st_size // 1024
        print(f"    · {f.name:<45} {size_kb:>5} KB")

    print()
    if len(pdfs) == len(docx_paths):
        print("  ✅ Todos los documentos generados correctamente.")
    elif len(pdfs) > 0:
        print(f"  ⚠️  {len(pdfs)}/{len(docx_paths)} PDF generados. Revisá los errores arriba.")
    else:
        print("  ⚠️  PDFs no generados. Abrí los .docx en Word y guardá como PDF manualmente.")
    print("═" * 60)
    print()


if __name__ == "__main__":
    main()
