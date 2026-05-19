"""
Genera Portal_B2B_P2_Playbook_v3.docx — v3 agrega la sección completa
de "Integración con Telegram" sobre la base del v2.
Ejecutar: python generar_playbook_v3.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Portal_B2B_P2_Playbook_v3.docx")

# ─── Colores ──────────────────────────────────────────────────────────────────
C_DARK      = RGBColor(0x09, 0x09, 0x0B)   # zinc-950
C_ZINC800   = RGBColor(0x27, 0x27, 0x2A)   # zinc-800
C_ZINC600   = RGBColor(0x52, 0x52, 0x56)   # zinc-600
C_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
C_ORANGE    = RGBColor(0xF9, 0x73, 0x16)   # orange-500
C_SKY       = RGBColor(0x0E, 0xA5, 0xE9)   # sky-500
C_EMERALD   = RGBColor(0x10, 0xB9, 0x81)   # emerald-500
C_TEAL      = RGBColor(0x14, 0xB8, 0xA6)   # teal-500
C_AMBER     = RGBColor(0xF5, 0x9E, 0x0B)   # amber-500
C_RED       = RGBColor(0xEF, 0x44, 0x44)   # red-500
C_VIOLET    = RGBColor(0x8B, 0x5C, 0xF6)   # violet-500
C_ZINC400   = RGBColor(0xA1, 0xA1, 0xAA)   # zinc-400 (cancelado)
C_ROSE      = RGBColor(0xF4, 0x3F, 0x5E)   # rose-500

# Estado → color de fondo (light) para celdas de tabla
STATUS_COLORS = {
    "pendiente":        RGBColor(0x16, 0x2A, 0x1A),
    "en_revision":      RGBColor(0x0C, 0x1A, 0x2E),
    "cotizado":         RGBColor(0x2A, 0x22, 0x06),
    "aprobado_parcial": RGBColor(0x2A, 0x18, 0x06),
    "aprobado":         RGBColor(0x06, 0x24, 0x18),
    "pagado":           RGBColor(0x1E, 0x10, 0x40),
    "rechazado":        RGBColor(0x2A, 0x08, 0x08),
    "cerrado":          RGBColor(0x0A, 0x14, 0x2A),
    "cerrado_pagado":   RGBColor(0x04, 0x20, 0x1E),
    "en_conflicto":     RGBColor(0x2A, 0x06, 0x06),
    "cancelado":        RGBColor(0x18, 0x18, 0x1C),
}

STATUS_TEXT_COLORS = {
    "pendiente":        C_EMERALD,
    "en_revision":      C_SKY,
    "cotizado":         C_AMBER,
    "aprobado_parcial": C_ORANGE,
    "aprobado":         C_EMERALD,
    "pagado":           C_VIOLET,
    "rechazado":        C_RED,
    "cerrado":          C_SKY,
    "cerrado_pagado":   C_TEAL,
    "en_conflicto":     C_RED,
    "cancelado":        C_ZINC400,
}

STATUS_LABELS = {
    "pendiente":        "Pendiente",
    "en_revision":      "En revisión",
    "cotizado":         "Cotizado",
    "aprobado_parcial": "Aprobado parcial",
    "aprobado":         "Aprobado",
    "pagado":           "Pagado · Por entregar",
    "rechazado":        "Rechazado",
    "cerrado":          "Cerrado",
    "cerrado_pagado":   "Cerrado · Pagado",
    "en_conflicto":     "En conflicto",
    "cancelado":        "Cancelado",
}

# ─── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}')
    tcPr.append(shd)

def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

def add_heading(doc, text, level=1, color=None, size=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size or (22 if level == 1 else 15 if level == 2 else 12))
    run.font.color.rgb = color or (C_ORANGE if level == 1 else C_WHITE if level == 2 else C_SKY)
    return p

def add_body(doc, text, color=None, bold=False, size=10, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(0.7)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color or C_ZINC400
    run.bold = bold
    return p

def add_bullet(doc, text, color=None, indent_level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Cm(0.5 + indent_level * 0.7)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = color or C_ZINC400

def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run('─' * 80)
    run.font.size = Pt(7)
    run.font.color.rgb = C_ZINC800

def add_page_break(doc):
    doc.add_page_break()

def add_status_badge_row(doc, statuses_list):
    """Dibuja badges de estados como tabla de 1 fila."""
    table = doc.add_table(rows=1, cols=len(statuses_list))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, status in enumerate(statuses_list):
        cell = table.rows[0].cells[i]
        cell.width = Cm(3.2)
        set_cell_bg(cell, STATUS_COLORS[status])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(STATUS_LABELS[status])
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.color.rgb = STATUS_TEXT_COLORS[status]
    doc.add_paragraph()

def make_flow_table(doc, rows_data):
    """Tabla de dos columnas: PASO | DESCRIPCION."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0]
    for i, label in enumerate(['PASO / ESTADO', 'DESCRIPCION Y ACCIONES DISPONIBLES']):
        hdr.cells[i].width = Cm(5.5 if i == 0 else 13)
        set_cell_bg(hdr.cells[i], C_ZINC800)
        p = hdr.cells[i].paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(label)
        run.font.size = Pt(8.5)
        run.font.bold = True
        run.font.color.rgb = C_WHITE

    for (status_key, desc_lines) in rows_data:
        row = table.add_row()
        # Col 0: estado
        c0 = row.cells[0]
        c0.width = Cm(5.5)
        set_cell_bg(c0, STATUS_COLORS.get(status_key, C_ZINC800))
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p0.paragraph_format.space_before = Pt(6)
        p0.paragraph_format.space_after  = Pt(6)
        run0 = p0.add_run(STATUS_LABELS.get(status_key, status_key.upper()))
        run0.font.size = Pt(9)
        run0.font.bold = True
        run0.font.color.rgb = STATUS_TEXT_COLORS.get(status_key, C_WHITE)

        # Col 1: descripción
        c1 = row.cells[1]
        c1.width = Cm(13)
        set_cell_bg(c1, C_DARK)
        for j, line in enumerate(desc_lines):
            if j == 0:
                p1 = c1.paragraphs[0]
            else:
                p1 = c1.add_paragraph()
            p1.paragraph_format.space_before = Pt(2)
            p1.paragraph_format.space_after  = Pt(2)
            is_bullet = line.startswith('  •') or line.startswith('  →') or line.startswith('    ')
            indent = Cm(0.4) if is_bullet else Cm(0.2)
            p1.paragraph_format.left_indent = indent
            run1 = p1.add_run(line)
            run1.font.size = Pt(9)
            if line.startswith('  →') or line.startswith('    →'):
                run1.font.color.rgb = C_EMERALD
                run1.font.italic = True
            elif line.startswith('  •'):
                run1.font.color.rgb = C_ZINC400
            elif '⚡' in line or '💬' in line or '🏦' in line or '💵' in line:
                run1.font.color.rgb = C_AMBER
            else:
                run1.font.color.rgb = C_ZINC400
    doc.add_paragraph()

# ─── Documento ────────────────────────────────────────────────────────────────

doc = Document()

# Configurar márgenes
for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)

# Fondo oscuro general no es posible en DOCX; se trabaja con secciones de texto.

# ════════════════════════════════════════════════════
# PORTADA
# ════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
r = p.add_run('PORTAL B2B AUTOPARTES')
r.font.size = Pt(28)
r.font.bold = True
r.font.color.rgb = C_ORANGE

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Playbook de Flujos y Máquina de Estados — v3.0')
r2.font.size = Pt(14)
r2.font.color.rgb = C_ZINC400

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('Flujo completo + Integración Telegram · Mayo 2026')
r3.font.size = Pt(10)
r3.font.color.rgb = C_ZINC600

add_separator(doc)

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run(
    'Este documento describe el flujo FINAL de pedidos del Portal B2B incluyendo\n'
    'el nuevo flujo de cobro/entrega, modales de aprobacion, WhatsApp y resolucion de conflictos.'
)
r4.font.size = Pt(10)
r4.font.color.rgb = C_ZINC400
r4.font.italic = True

add_page_break(doc)

# ════════════════════════════════════════════════════
# SECCION 1 — ROLES
# ════════════════════════════════════════════════════
add_heading(doc, '1. ROLES DEL SISTEMA', level=1)

roles = [
    ('🏭  TALLER',    C_EMERALD,  [
        'Crea pedidos de repuestos para su vehiculo.',
        'Ve las cotizaciones del vendedor y las aprueba, rechaza o aprueba parcialmente.',
        'Al aprobar, debe elegir metodo de pago (Transferencia / Efectivo).',
        'Puede avisar al vendedor por WhatsApp con mensaje pre-cargado.',
        'Puede iniciar reclamos sobre pedidos ya cerrados.',
        'Ve el timeline con descripciones especificas para su rol.',
    ]),
    ('💼  VENDEDOR',  C_ORANGE, [
        'Toma pedidos de la cola general o los que le asignaron.',
        'Marca en revision, arma y envia la cotizacion.',
        'Cuando el taller aprueba, puede: Marcar como Pagado (recibio el dinero) o Entregado y Pagado (todo en un paso).',
        'Si el pedido esta en "Pagado", marca la entrega cuando entrega los repuestos.',
        'Puede resolver conflictos activos (ver Tarea 4).',
        'Ve el timeline con descripciones especificas para vendedor.',
    ]),
    ('🔐  ADMIN',     C_VIOLET, [
        'Supervision completa de todos los pedidos y talleres.',
        'Puede confirmar pago de pedidos cerrados (cerrado → cerrado_pagado).',
        'Puede resolver conflictos igual que el vendedor.',
        'Ve el ranking de vendedores y metricas globales.',
        'Puede gestionar usuarios y asignar vendedores a talleres.',
    ]),
]

for title, color, bullets in roles:
    p = add_heading(doc, title, level=2, color=color)
    for b in bullets:
        add_bullet(doc, b)
    doc.add_paragraph()

add_page_break(doc)

# ════════════════════════════════════════════════════
# SECCION 2 — MAQUINA DE ESTADOS COMPLETA
# ════════════════════════════════════════════════════
add_heading(doc, '2. MAQUINA DE ESTADOS COMPLETA', level=1)

add_body(doc, 'El sistema maneja 11 estados posibles para un pedido. Cada transicion genera un evento en el historial con fecha, actor y comentario.', size=10)
doc.add_paragraph()

# Tabla de todos los estados
add_status_badge_row(doc, ['pendiente', 'en_revision', 'cotizado', 'aprobado_parcial', 'aprobado', 'pagado'])
add_status_badge_row(doc, ['rechazado', 'cerrado', 'cerrado_pagado', 'en_conflicto', 'cancelado'])

add_heading(doc, 'Diagrama de transiciones', level=2)

transitions = [
    ('pendiente',        'Nuevo pedido creado por el taller'),
    ('en_revision',      'Vendedor toma y marca en revision'),
    ('cotizado',         'Vendedor envia la cotizacion'),
    ('aprobado',         'Taller aprueba todos los items'),
    ('aprobado_parcial', 'Taller aprueba solo algunos items'),
    ('rechazado',        'Taller rechaza la cotizacion completa'),
    ('pagado',           'Vendedor registra el pago recibido (mercaderia pendiente de entrega)'),
    ('cerrado',          'Vendedor cierra (caso legacy) o resolucio de conflicto sin cancelacion'),
    ('cerrado_pagado',   'Vendedor marca entrega + pago, o Admin confirma pago'),
    ('en_conflicto',     'Taller inicia reclamo desde estado cerrado'),
    ('cancelado',        'Conflicto resuelto con resultado "Cancelado"'),
]

table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, lbl in enumerate(['ESTADO', 'DESCRIPCION', 'QUIEN LO ACTIVA']):
    set_cell_bg(table.rows[0].cells[i], C_ZINC800)
    p = table.rows[0].cells[i].paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(lbl)
    r.font.size = Pt(8.5)
    r.font.bold = True
    r.font.color.rgb = C_WHITE

actores = {
    'pendiente':        'Taller',
    'en_revision':      'Vendedor',
    'cotizado':         'Vendedor',
    'aprobado':         'Taller',
    'aprobado_parcial': 'Taller',
    'rechazado':        'Taller',
    'pagado':           'Vendedor',
    'cerrado':          'Vendedor / Sistema',
    'cerrado_pagado':   'Vendedor / Admin',
    'en_conflicto':     'Taller',
    'cancelado':        'Admin / Vendedor',
}

for (status, desc) in transitions:
    row = table.add_row()
    set_cell_bg(row.cells[0], STATUS_COLORS[status])
    p0 = row.cells[0].paragraphs[0]
    p0.paragraph_format.space_before = Pt(5)
    p0.paragraph_format.space_after  = Pt(5)
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run(STATUS_LABELS[status])
    r0.font.size = Pt(8.5)
    r0.font.bold = True
    r0.font.color.rgb = STATUS_TEXT_COLORS[status]

    set_cell_bg(row.cells[1], C_DARK)
    p1 = row.cells[1].paragraphs[0]
    p1.paragraph_format.space_before = Pt(5)
    p1.paragraph_format.space_after  = Pt(5)
    p1.paragraph_format.left_indent  = Cm(0.2)
    r1 = p1.add_run(desc)
    r1.font.size = Pt(9)
    r1.font.color.rgb = C_ZINC400

    set_cell_bg(row.cells[2], C_DARK)
    p2 = row.cells[2].paragraphs[0]
    p2.paragraph_format.space_before = Pt(5)
    p2.paragraph_format.space_after  = Pt(5)
    p2.paragraph_format.left_indent  = Cm(0.2)
    r2 = p2.add_run(actores[status])
    r2.font.size = Pt(9)
    r2.font.color.rgb = C_ORANGE

doc.add_paragraph()

add_page_break(doc)

# ════════════════════════════════════════════════════
# SECCION 3 — FLUJO PRINCIPAL PASO A PASO
# ════════════════════════════════════════════════════
add_heading(doc, '3. FLUJO PRINCIPAL — PASO A PASO', level=1)
add_body(doc, 'Este es el flujo estandar de un pedido desde su creacion hasta el cierre, para el caso mas comun (aprobacion total).', size=10)
doc.add_paragraph()

main_flow = [
    ('pendiente', [
        'Taller crea el pedido: vehiculo, items (repuesto, calidad, cantidad, fotos).',
        '  • El pedido entra a la cola general y todos los vendedores pueden verlo.',
        '  • El taller puede eliminar el pedido mientras esta pendiente.',
        '  → Siguiente accion: vendedor toma el pedido.',
    ]),
    ('en_revision', [
        'Vendedor toma el pedido (self-assign) y lo marca en revision.',
        '  • El pedido queda asignado a ese vendedor.',
        '  • El vendedor puede liberarlo de vuelta a la cola si no puede atenderlo.',
        '  • Puede armar la cotizacion desde este estado.',
        '  → Siguiente accion: vendedor envia cotizacion.',
    ]),
    ('cotizado', [
        'Vendedor arma y envia la cotizacion con precios, cantidad ofrecida y fotos.',
        '  • El taller recibe la cotizacion y puede comparar lo pedido vs lo cotizado.',
        '  • Taller ve: precio unitario, cantidad ofrecida vs solicitada, subtotales.',
        '  → El taller tiene 3 opciones: Aprobar Todo, Aprobacion Parcial o Rechazar.',
    ]),
    ('aprobado', [
        'Taller aprueba toda la cotizacion.',
        '  • ANTES de aprobar: se muestra un modal "Como vas a pagar?"',
        '  • Opciones de pago: Transferencia (default) o Efectivo.',
        '  • Al confirmar, el pedido pasa a Aprobado.',
        '  ⚡ POST-APROBACION: aparece boton WhatsApp con mensaje pre-cargado:',
        '     "Hola [Vendedor], soy el taller [Nombre]. Te apruebo el pedido',
        '      #PED-XXXX por un total de $XX.XXX. Lo voy a pagar por transferencia.',
        '      Avisame para coordinar."',
        '  → Siguiente accion: vendedor gestiona el cobro y la entrega.',
    ]),
    ('aprobado_parcial', [
        'Taller aprueba solo algunos items y rechaza los demas.',
        '  • Mismo modal de metodo de pago y WhatsApp que en aprobacion total.',
        '  • El total es la suma solo de los items aprobados.',
        '  → Misma gestion de cobro/entrega que estado Aprobado.',
    ]),
    ('pagado', [
        'Vendedor registra que el taller realizo el pago (pero no entrego aun).',
        '  • Estado intermedio: pago recibido, mercaderia pendiente de entrega.',
        '  • El taller ve: "Pago registrado. Los repuestos estan siendo preparados."',
        '  • El vendedor ve: "Registraste el pago. Cuando entregues, marca la entrega."',
        '  → Siguiente accion: vendedor entrega los repuestos.',
    ]),
    ('cerrado_pagado', [
        'Pago confirmado y mercaderia entregada. Pedido completado.',
        '  • Se activa desde "Pagado" con el boton "Marcar como Entregado".',
        '  • O directamente desde "Aprobado/Aprobado parcial" con "Entregado y Pagado".',
        '  • Tambien lo puede confirmar el Admin desde el estado "Cerrado".',
        '  → Estado final. No hay mas acciones posibles.',
    ]),
]

make_flow_table(doc, main_flow)

add_page_break(doc)

# ════════════════════════════════════════════════════
# SECCION 4 — FLUJO DE COBRO Y ENTREGA (DETALLE)
# ════════════════════════════════════════════════════
add_heading(doc, '4. FLUJO DE COBRO Y ENTREGA — DETALLE COMPLETO', level=1)
add_body(doc, 'Una vez que el taller aprueba (total o parcial), el vendedor tiene tres caminos para cerrar el pedido.', size=10)
doc.add_paragraph()

add_heading(doc, 'Camino A — Flujo rapido (pago y entrega al mismo tiempo)', level=2)
add_bullet(doc, 'Estado: aprobado o aprobado_parcial')
add_bullet(doc, 'Boton: "Entregado y Pagado" (disponible en pagina de detalle y en el Drawer lateral)')
add_bullet(doc, 'Resultado: status → cerrado_pagado + evento pedido_entregado')
add_bullet(doc, 'Uso: cuando el taller paga al retirar los repuestos (pago en mano + entrega inmediata)')
doc.add_paragraph()

add_heading(doc, 'Camino B — Flujo en dos pasos (pago primero, entrega despues)', level=2)
add_bullet(doc, 'PASO 1 — Estado: aprobado o aprobado_parcial')
add_bullet(doc,   'Boton: "Marcar como Pagado"', indent_level=1)
add_bullet(doc,   'Resultado: status → pagado + evento pedido_marcado_pagado', indent_level=1)
add_bullet(doc,   'Uso: el taller hizo una transferencia pero los repuestos se entregan despues', indent_level=1)
doc.add_paragraph()
add_bullet(doc, 'PASO 2 — Estado: pagado')
add_bullet(doc,   'Boton: "Marcar como Entregado"', indent_level=1)
add_bullet(doc,   'Resultado: status → cerrado_pagado + evento pedido_entregado', indent_level=1)
add_bullet(doc,   'Uso: cuando se va a hacer la entrega efectiva de la mercaderia', indent_level=1)
doc.add_paragraph()

add_heading(doc, 'Camino C — Confirmacion por Admin (caso legacy)', level=2)
add_bullet(doc, 'Si el pedido llego al estado "cerrado" por flujo antiguo, el Admin puede confirmar pago.')
add_bullet(doc, 'Boton "Confirmar pago" en el Drawer (solo para Admin).')
add_bullet(doc, 'Resultado: status → cerrado_pagado + evento pedido_pagado')
doc.add_paragraph()

# Diagrama simplificado de cobro
add_heading(doc, 'Diagrama — estados de cobro/entrega', level=2)

cobro_flow = [
    ('aprobado', [
        'Estado post-aprobacion del taller.',
        '  • Boton A: "Entregado y Pagado" → cerrado_pagado (camino rapido)',
        '  • Boton B: "Marcar como Pagado" → pagado (camino en dos pasos)',
    ]),
    ('aprobado_parcial', [
        'Igual que Aprobado pero solo con items seleccionados.',
        '  • Mismo Boton A y Boton B disponibles.',
    ]),
    ('pagado', [
        'Pago registrado por el vendedor. Mercaderia pendiente de entrega.',
        '  • Boton: "Marcar como Entregado" → cerrado_pagado',
    ]),
    ('cerrado_pagado', [
        'Estado FINAL. Pago confirmado + mercaderia entregada.',
        '  • No hay mas acciones disponibles.',
        '  • Se contabiliza en las metricas de facturacion del vendedor.',
    ]),
]

make_flow_table(doc, cobro_flow)

add_page_break(doc)

# ════════════════════════════════════════════════════
# SECCION 5 — FLUJO DE APROBACION (MODAL + WHATSAPP)
# ════════════════════════════════════════════════════
add_heading(doc, '5. FLUJO DE APROBACION DEL TALLER — MODAL + WHATSAPP', level=1)
add_body(doc, 'Cuando el taller hace clic en "Aprobar Todo", el sistema no aprueba directamente. Primero muestra un modal de confirmacion.', size=10)
doc.add_paragraph()

add_heading(doc, 'Modal de metodo de pago', level=2)
add_bullet(doc, 'Titulo del modal: "Como vas a pagar este pedido?"')
add_bullet(doc, 'Subtitulo: "Selecciona el metodo de pago para informarle al vendedor."')
add_bullet(doc, 'Opciones (botones toggle exclusivos):')
add_bullet(doc,   'Transferencia bancaria (seleccionada por defecto)', indent_level=1)
add_bullet(doc,   'Efectivo', indent_level=1)
add_bullet(doc, 'Botones del modal:')
add_bullet(doc,   'Cancelar → cierra el modal sin cambios', indent_level=1)
add_bullet(doc,   'Confirmar aprobacion → aprueba el pedido y cierra el modal', indent_level=1)
doc.add_paragraph()

add_heading(doc, 'Boton WhatsApp post-aprobacion', level=2)
add_body(doc, 'Inmediatamente despues de la aprobacion, aparece una barra destacada con el boton de WhatsApp.', size=9.5)
doc.add_paragraph()
add_bullet(doc, 'El toggle de metodo de pago sigue visible para cambiar antes de enviar.')
add_bullet(doc, 'El boton abre WhatsApp (web o app) con este mensaje pre-cargado:')
doc.add_paragraph()

# Cuadro con el mensaje WA
t = doc.add_table(rows=1, cols=1)
t.alignment = WD_TABLE_ALIGNMENT.LEFT
set_cell_bg(t.rows[0].cells[0], RGBColor(0x06, 0x24, 0x18))
p = t.rows[0].cells[0].paragraphs[0]
p.paragraph_format.left_indent  = Cm(0.5)
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after  = Pt(8)
r = p.add_run(
    '"Hola [Nombre del Vendedor], soy el taller [Nombre del Taller].\n'
    'Te apruebo el pedido #PED-XXXX por un total de $XX.XXX.\n'
    'Lo voy a pagar por [transferencia bancaria / efectivo].\n'
    'Avisame para coordinar."'
)
r.font.size = Pt(9.5)
r.font.color.rgb = C_EMERALD
r.font.italic = True
doc.add_paragraph()

add_body(doc, 'Variables dinamicas del mensaje:', bold=True, size=9.5)
add_bullet(doc, '[Nombre del Vendedor] → nombre del vendedor asignado al pedido (order.assignedVendorName)')
add_bullet(doc, '[Nombre del Taller] → nombre del taller del usuario autenticado (user.workshopName)')
add_bullet(doc, '#PED-XXXX → numero de pedido formateado (workshopOrderNumber)')
add_bullet(doc, '$XX.XXX → total calculado de los items aprobados (formatCurrency)')
add_bullet(doc, '[metodo de pago] → el que selecciono el taller en el modal')
doc.add_paragraph()

add_body(doc, 'Nota: el boton de WhatsApp tambien aparece para pedidos ya aprobados al recargar la pagina, con el metodo de pago por defecto (Transferencia). El taller puede cambiar el toggle antes de enviar.', size=9, color=C_ZINC600)
doc.add_paragraph()

add_heading(doc, 'Flujo de aprobacion parcial', level=2)
add_bullet(doc, 'El taller selecciona los items a aprobar con checkboxes en la comparativa de cotizacion.')
add_bullet(doc, 'Al hacer clic en "Seleccionar (N)", se muestra el MISMO modal de metodo de pago.')
add_bullet(doc, 'Post-aprobacion parcial: mismo boton WhatsApp con el total de items aprobados.')
doc.add_paragraph()

add_page_break(doc)

# ════════════════════════════════════════════════════
# SECCION 6 — FLUJO DE CONFLICTOS Y RESOLUCIONES
# ════════════════════════════════════════════════════
add_heading(doc, '6. FLUJO DE CONFLICTOS Y RESOLUCIONES', level=1)
add_body(doc, 'Cuando hay un problema despues de cerrar un pedido, el taller puede iniciar un reclamo. El flujo tiene 3 fases.', size=10)
doc.add_paragraph()

add_heading(doc, 'FASE 1 — Taller inicia el reclamo', level=2)

conflict_flow_taller = [
    ('cerrado', [
        'El pedido esta cerrado (entregado pero no pagado, o cerrado por flujo antiguo).',
        '  • El taller ve el boton: "Iniciar reclamo" (zona de alerta amarilla).',
        '  • Al hacer clic, se abre un modal con textarea obligatoria.',
        '  • El taller describe el problema: repuesto incorrecto, faltante, etc.',
        '  → Al confirmar: status → en_conflicto + evento reclamo_iniciado',
    ]),
    ('en_conflicto', [
        'VISTA DEL TALLER en estado en_conflicto:',
        '  • Mensaje fijo: "Pedido en conflicto. Comunicate por WhatsApp para coordinar."',
        '  • Boton WhatsApp con mensaje pre-cargado:',
        '  ⚡ "Hola, soy [Taller]. Tengo un reclamo en el pedido #PED-XXXX.",',
        '     "Podemos coordinar la resolucion?"',
        '  • No hay mas acciones disponibles para el taller.',
    ]),
]

make_flow_table(doc, conflict_flow_taller)

add_heading(doc, 'FASE 2 — Vendedor/Admin gestiona el conflicto', level=2)
add_bullet(doc, 'El vendedor ve una alerta roja urgente en la pagina del pedido y en el Drawer lateral.')
add_bullet(doc, 'El Drawer del Admin muestra boton "Gestionar conflicto" que lleva al detalle.')
add_bullet(doc, 'Ambos ven el historial con el motivo del reclamo del taller.')
add_bullet(doc, 'Boton disponible: "Conflicto Solucionado" (vendedor y admin).')
doc.add_paragraph()

add_heading(doc, 'FASE 3 — Registro de la resolucion (modal)', level=2)
add_body(doc, 'Al hacer clic en "Conflicto Solucionado", se abre el modal de resolucion:', size=9.5)
doc.add_paragraph()
add_bullet(doc, 'Campo 1: Select "Tipo de resolucion" (obligatorio)')
add_bullet(doc,   'Devolucion total', indent_level=1)
add_bullet(doc,   'Devolucion parcial', indent_level=1)
add_bullet(doc,   'Se aplico descuento', indent_level=1)
add_bullet(doc,   'Resuelto sin cambios', indent_level=1)
add_bullet(doc,   'Cancelado', indent_level=1)
add_bullet(doc, 'Campo 2: Textarea "Detalle del acuerdo" (obligatorio)')
add_bullet(doc,   'Descripcion libre del acuerdo alcanzado', indent_level=1)
add_bullet(doc, 'Si se elige "Cancelado": aparece alerta de confirmacion adicional.')
doc.add_paragraph()

add_heading(doc, 'Resultado de la resolucion', level=2)

resolution_flow = [
    ('cerrado', [
        'Se activa cuando la resolucion es: Devolucion total, Devolucion parcial,',
        'Se aplico descuento, o Resuelto sin cambios.',
        '  → Evento insertado: conflicto_resuelto con "Resolucion: [tipo]. [detalle]"',
    ]),
    ('cancelado', [
        'Se activa cuando la resolucion es: Cancelado.',
        '  • El taller ve: "Este pedido fue cancelado como resolucion al conflicto."',
        '  • El vendedor ve: "Pedido cancelado como parte de la resolucion."',
        '  → Evento insertado: conflicto_resuelto con detalle del acuerdo.',
        '  → Estado FINAL. No hay mas acciones posibles.',
    ]),
]

make_flow_table(doc, resolution_flow)

add_page_break(doc)

# ════════════════════════════════════════════════════
# SECCION 7 — TIMELINE INTELIGENTE POR ROL
# ════════════════════════════════════════════════════
add_heading(doc, '7. TIMELINE INTELIGENTE — DESCRIPCIONES POR ROL', level=1)
add_body(doc, 'El tracker de estados (barra de progreso del pedido) muestra descripciones distintas segun quien esta mirando.', size=10)
doc.add_paragraph()

# Tabla de descripciones por rol
table = doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.LEFT

headers = ['ESTADO', 'TALLER ve', 'VENDEDOR ve', 'ADMIN ve']
hcolors = [C_WHITE, C_EMERALD, C_ORANGE, C_VIOLET]
for i, (lbl, col) in enumerate(zip(headers, hcolors)):
    set_cell_bg(table.rows[0].cells[i], C_ZINC800)
    p = table.rows[0].cells[i].paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(lbl)
    r.font.size = Pt(8.5)
    r.font.bold = True
    r.font.color.rgb = col

role_desc = [
    ('Pendiente', [
        'Tu pedido fue enviado. Un vendedor lo revisara pronto.',
        'Nuevo pedido en la cola. Tomalo para gestionar.',
        'Pedido en espera de ser tomado por un vendedor.',
    ]),
    ('En revision', [
        'El vendedor esta consultando disponibilidad y precios.',
        'Consulta disponibilidad y arma la cotizacion para el taller.',
        'El vendedor esta consultando disponibilidad y armando la cotizacion.',
    ]),
    ('Cotizado', [
        'Cotizacion recibida. Revisa los items y aproba para continuar.',
        'Cotizacion enviada. Esperando respuesta del taller.',
        'Cotizacion enviada al taller. Pendiente de aprobacion.',
    ]),
    ('Resolucion', [
        'Tu respuesta fue registrada. El vendedor preparara lo acordado.',
        'Respuesta del taller recibida. Coordina la preparacion y entrega.',
        'El taller respondio la cotizacion.',
    ]),
    ('Cerrado', [
        'Pedido completado. Los repuestos fueron entregados.',
        'Pedido completado y entregado al taller.',
        'Pedido finalizado.',
    ]),
]

for (step, descs) in role_desc:
    row = table.add_row()
    set_cell_bg(row.cells[0], C_ZINC800)
    p0 = row.cells[0].paragraphs[0]
    p0.paragraph_format.space_before = Pt(5)
    p0.paragraph_format.space_after  = Pt(5)
    p0.paragraph_format.left_indent  = Cm(0.2)
    r0 = p0.add_run(step)
    r0.font.size = Pt(9)
    r0.font.bold = True
    r0.font.color.rgb = C_WHITE

    colors = [C_EMERALD, C_ORANGE, C_VIOLET]
    for j, (desc, color) in enumerate(zip(descs, colors)):
        set_cell_bg(row.cells[j+1], C_DARK)
        p = row.cells[j+1].paragraphs[0]
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after  = Pt(5)
        p.paragraph_format.left_indent  = Cm(0.2)
        r = p.add_run(desc)
        r.font.size = Pt(8.5)
        r.font.color.rgb = color
        r.font.italic = True

doc.add_paragraph()

add_body(doc, 'Ademas, el timeline muestra el NOMBRE REAL del actor que cambio cada estado (resuelto desde la base de datos, no un ID). Ejemplo: "Maria Garcia * Aprobado".', size=9.5, color=C_ZINC400)
doc.add_paragraph()

add_page_break(doc)

# ════════════════════════════════════════════════════
# SECCION 8 — FLUJOS ESPECIALES Y CASOS BORDE
# ════════════════════════════════════════════════════
add_heading(doc, '8. FLUJOS ESPECIALES Y CASOS BORDE', level=1)

add_heading(doc, '8.1 Pedido rechazado', level=2)
add_bullet(doc, 'El taller rechaza la cotizacion completa.')
add_bullet(doc, 'Status: rechazado. Es un estado terminal.')
add_bullet(doc, 'El vendedor puede ver el rechazo pero no puede re-cotizar (debe crear un nuevo pedido).')
add_bullet(doc, 'El taller puede ver el motivo del rechazo en el historial.')
doc.add_paragraph()

add_heading(doc, '8.2 Pedido sin cotizacion disponible', level=2)
add_bullet(doc, 'Si el vendedor no tiene stock de ningun item, puede marcar todos como "Sin stock".')
add_bullet(doc, 'La cotizacion enviada tendra 0 items. El taller ve el mensaje: "El vendedor no tiene stock para ningun item."')
doc.add_paragraph()

add_heading(doc, '8.3 Eliminacion de pedidos', level=2)
add_bullet(doc, 'El TALLER puede eliminar un pedido en estado "pendiente" desde el detalle.')
add_bullet(doc, 'El VENDEDOR puede eliminar pedidos en "pendiente" o "en_revision".')
add_bullet(doc, 'El ADMIN puede ver pedidos eliminados (soft-delete; campo deletedAt).')
add_bullet(doc, 'Los pedidos eliminados no aparecen en la lista del taller.')
doc.add_paragraph()

add_heading(doc, '8.4 Reasignacion de vendedor', level=2)
add_bullet(doc, 'Solo el ADMIN puede reasignar un pedido a otro vendedor.')
add_bullet(doc, 'Se hace desde el panel de admin (no desde el detalle del pedido).')
doc.add_paragraph()

add_heading(doc, '8.5 Cola general de pedidos', level=2)
add_bullet(doc, 'Pedidos sin vendedor asignado son visibles en la cola para todos los vendedores.')
add_bullet(doc, 'El primer vendedor en hacer clic en "Tomar pedido" se lo asigna (self-assign).')
add_bullet(doc, 'El vendedor puede "liberar" el pedido de vuelta a la cola desde el estado pendiente o en_revision.')
doc.add_paragraph()

add_page_break(doc)

# ════════════════════════════════════════════════════
# SECCION 9 — EVENTOS DEL HISTORIAL
# ════════════════════════════════════════════════════
add_heading(doc, '9. EVENTOS DEL HISTORIAL (order_events)', level=1)
add_body(doc, 'Cada cambio de estado genera un evento persistente con: actor, accion, comentario y timestamp. El timeline visible en el portal se construye desde esta tabla.', size=10)
doc.add_paragraph()

events = [
    ('pedido_creado',               '📋', 'Taller',          'Se crea el pedido con todos sus items.'),
    ('pedido_en_revision',          '🔍', 'Vendedor',         'Vendedor marca el pedido como en revision.'),
    ('pedido_tomado',               '🙋', 'Vendedor',         'Vendedor toma el pedido de la cola general.'),
    ('pedido_liberado',             '🔓', 'Vendedor',         'Vendedor libera el pedido de vuelta a la cola.'),
    ('cotizacion_enviada',          '📤', 'Vendedor',         'Cotizacion enviada con todos sus items y precios.'),
    ('cotizacion_aprobada',         '✅', 'Taller',           'Taller aprueba la cotizacion completa.'),
    ('cotizacion_aprobada_parcial', '⚡', 'Taller',           'Taller aprueba parcialmente la cotizacion.'),
    ('cotizacion_rechazada',        '❌', 'Taller',           'Taller rechaza la cotizacion completa.'),
    ('pedido_marcado_pagado',       '💰', 'Vendedor',         'Vendedor registra el pago (→ estado pagado).'),
    ('pedido_entregado',            '📦', 'Vendedor',         'Vendedor confirma entrega de mercaderia (→ cerrado_pagado).'),
    ('pedido_cerrado',              '🔒', 'Vendedor',         'Cierre del pedido (flujo legacy → cerrado).'),
    ('pedido_pagado',               '💳', 'Admin',            'Admin confirma pago (cerrado → cerrado_pagado).'),
    ('reclamo_iniciado',            '⚠️', 'Taller',           'Taller inicia un reclamo (cerrado → en_conflicto).'),
    ('conflicto_resuelto',          '🤝', 'Admin / Vendedor', 'Conflicto resuelto con tipo y detalle del acuerdo.'),
    ('comentario',                  '💬', 'Cualquiera',       'Comentario libre en el historial del pedido.'),
]

table = doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, lbl in enumerate(['ICONO', 'ACCION', 'ACTOR', 'DESCRIPCION']):
    set_cell_bg(table.rows[0].cells[i], C_ZINC800)
    p = table.rows[0].cells[i].paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(lbl)
    r.font.size = Pt(8.5)
    r.font.bold = True
    r.font.color.rgb = C_WHITE

for (action, icon, actor, desc) in events:
    row = table.add_row()
    set_cell_bg(row.cells[0], C_DARK)
    p0 = row.cells[0].paragraphs[0]
    p0.paragraph_format.space_before = Pt(4)
    p0.paragraph_format.space_after  = Pt(4)
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run(icon)
    r0.font.size = Pt(10)

    set_cell_bg(row.cells[1], C_DARK)
    p1 = row.cells[1].paragraphs[0]
    p1.paragraph_format.space_before = Pt(4)
    p1.paragraph_format.space_after  = Pt(4)
    p1.paragraph_format.left_indent  = Cm(0.2)
    r1 = p1.add_run(action)
    r1.font.size = Pt(8.5)
    r1.font.color.rgb = C_SKY
    r1.font.bold = True

    set_cell_bg(row.cells[2], C_DARK)
    p2 = row.cells[2].paragraphs[0]
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after  = Pt(4)
    p2.paragraph_format.left_indent  = Cm(0.2)
    r2 = p2.add_run(actor)
    r2.font.size = Pt(8.5)
    r2.font.color.rgb = C_ORANGE

    set_cell_bg(row.cells[3], C_DARK)
    p3 = row.cells[3].paragraphs[0]
    p3.paragraph_format.space_before = Pt(4)
    p3.paragraph_format.space_after  = Pt(4)
    p3.paragraph_format.left_indent  = Cm(0.2)
    r3 = p3.add_run(desc)
    r3.font.size = Pt(8.5)
    r3.font.color.rgb = C_ZINC400

doc.add_paragraph()

add_page_break(doc)

# ════════════════════════════════════════════════════
# SECCION 10 — INTEGRACION CON TELEGRAM
# ════════════════════════════════════════════════════
add_heading(doc, '10. INTEGRACION CON TELEGRAM', level=1)
add_body(doc,
    'El Portal B2B notifica automaticamente al equipo de ventas en un grupo de Telegram '
    'y envia metricas privadas al administrador (Juan). Las notificaciones son en tiempo real, '
    'disparadas por Supabase via webhooks cada vez que cambia el estado de un pedido.', size=10)
doc.add_paragraph()

# ─── 10.1 Arquitectura ──────────────────────────────
add_heading(doc, '10.1 Arquitectura del sistema', level=2)

add_bullet(doc, 'Bot de Telegram: creado con @BotFather, expone API HTTP (sendMessage).')
add_bullet(doc, 'Supabase: dispara webhooks (HTTP POST) cuando hay INSERT/UPDATE en orders u order_events.')
add_bullet(doc, 'Vercel: aloja el endpoint /api/webhooks/supabase/telegram que recibe los webhooks.')
add_bullet(doc, 'Endpoint: valida el secret, resuelve contexto (workshop, vendor, monto) y dispara la API de Telegram.')
add_bullet(doc, 'CRM Web: el admin gestiona los @usuarios de Telegram de cada vendedor desde /admin/vendedores.')
doc.add_paragraph()

# Diagrama ASCII del flujo
t = doc.add_table(rows=1, cols=1)
t.alignment = WD_TABLE_ALIGNMENT.LEFT
set_cell_bg(t.rows[0].cells[0], RGBColor(0x0A, 0x14, 0x2E))
p = t.rows[0].cells[0].paragraphs[0]
p.paragraph_format.left_indent = Cm(0.5)
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after  = Pt(8)
r = p.add_run(
    'Taller / Vendedor (CRM)\n'
    '       |  cambia estado\n'
    '       v\n'
    'Supabase (orders + order_events)\n'
    '       |  webhook HTTP POST\n'
    '       v\n'
    'Vercel (/api/webhooks/supabase/telegram)\n'
    '       |  fetch sendMessage\n'
    '       v\n'
    'Telegram Bot API\n'
    '       |\n'
    '       +---> Grupo de ventas (todos)\n'
    '       +---> Chat privado de Juan (metricas)'
)
r.font.size = Pt(9)
r.font.name = 'Consolas'
r.font.color.rgb = C_SKY
doc.add_paragraph()

# ─── 10.2 Eventos del grupo ──────────────────────────
add_heading(doc, '10.2 Eventos que disparan notificaciones al GRUPO', level=2)

events_table = [
    ('🆕', 'NUEVO PEDIDO', 'pedido_creado', 'Taller crea un pedido nuevo'),
    ('🙋', 'TOMADO', 'pedido_tomado', 'Vendedor toma un pedido de la cola'),
    ('📝', 'COTIZADO', 'cotizacion_enviada', 'Vendedor envia la cotizacion'),
    ('🟢', 'APROBADO', 'cotizacion_aprobada', 'Taller aprueba (etiqueta al vendedor)'),
    ('🟡', 'APROBADO PARCIAL', 'cotizacion_aprobada_parcial', 'Taller aprueba algunos items (etiqueta vendedor)'),
    ('🔴', 'RECHAZADO', 'cotizacion_rechazada', 'Taller rechaza la cotizacion (etiqueta vendedor)'),
    ('💰', 'PAGO REGISTRADO', 'pedido_marcado_pagado', 'Vendedor confirma que recibio el pago'),
    ('📦', 'ENTREGADO Y COBRADO', 'pedido_entregado', 'Vendedor confirma entrega completa'),
    ('⚠️', 'CONFLICTO INICIADO', 'reclamo_iniciado', 'Taller inicia un reclamo (etiqueta vendedor)'),
    ('🤝', 'CONFLICTO RESUELTO', 'conflicto_resuelto', 'Admin/vendedor resuelve el conflicto'),
]

table = doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, lbl in enumerate(['ICONO', 'TAG VISIBLE', 'EVENTO', 'CUANDO SE DISPARA']):
    set_cell_bg(table.rows[0].cells[i], C_ZINC800)
    pp = table.rows[0].cells[i].paragraphs[0]
    pp.paragraph_format.space_before = Pt(4)
    pp.paragraph_format.space_after = Pt(4)
    rr = pp.add_run(lbl)
    rr.font.size = Pt(8.5)
    rr.font.bold = True
    rr.font.color.rgb = C_WHITE

for (icon, tag, event, desc) in events_table:
    row = table.add_row()
    for j, (val, color, bold) in enumerate([
        (icon, C_WHITE, False),
        (tag, C_AMBER, True),
        (event, C_SKY, True),
        (desc, C_ZINC400, False),
    ]):
        set_cell_bg(row.cells[j], C_DARK)
        pp = row.cells[j].paragraphs[0]
        pp.paragraph_format.space_before = Pt(4)
        pp.paragraph_format.space_after = Pt(4)
        pp.paragraph_format.left_indent = Cm(0.2)
        if j == 0: pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = pp.add_run(val)
        rr.font.size = Pt(9)
        rr.font.bold = bold
        rr.font.color.rgb = color

doc.add_paragraph()
add_body(doc, 'Eventos silenciados (no llegan al grupo): pedido_en_revision, pedido_liberado, pedido_cerrado, comentario.', size=9, color=C_ZINC600)
doc.add_paragraph()

# ─── 10.3 Etiquetas dinamicas ────────────────────────
add_heading(doc, '10.3 Etiquetas dinamicas de vendedores (@usuarios)', level=2)
add_body(doc,
    'Cuando el evento requiere la atencion de un vendedor especifico (aprobacion, rechazo, conflicto), '
    'el bot lo etiqueta por su @usuario de Telegram para que reciba la notificacion push.', size=10)
doc.add_paragraph()

add_bullet(doc, 'El @usuario se guarda en profiles.telegram_username (columna text nullable).')
add_bullet(doc, 'Lo configura el admin desde /admin/vendedores → seleccionar vendedor → Editar → campo "Usuario de Telegram".')
add_bullet(doc, 'Al guardar, el sistema normaliza: saca @, espacios y se queda solo con el username (ej: "Franco_San_Martin").')
add_bullet(doc, 'En tiempo real, el webhook consulta profiles.telegram_username del vendor_id asignado al pedido.')
add_bullet(doc, 'Si tiene username configurado: el mensaje incluye "@Franco_San_Martin" (genera ping en Telegram).')
add_bullet(doc, 'Si NO tiene username: fallback al nombre real en negrita (ej: <b>Lucas Pereyra</b>) sin ping.')
doc.add_paragraph()

# Ejemplo de mensaje
add_body(doc, 'EJEMPLO DE MENSAJE EN EL GRUPO (con etiqueta):', bold=True, size=9.5)
t = doc.add_table(rows=1, cols=1)
t.alignment = WD_TABLE_ALIGNMENT.LEFT
set_cell_bg(t.rows[0].cells[0], RGBColor(0x06, 0x24, 0x18))
p = t.rows[0].cells[0].paragraphs[0]
p.paragraph_format.left_indent = Cm(0.5)
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(6)
r = p.add_run(
    '🟢 [APROBADO]\n'
    '🏢 Taller: Carrocerias del Norte\n'
    '📦 Pedido: #PED-0109\n'
    '💰 Monto: $17.000\n'
    '\n'
    '🔔 @Franco_San_Martin, coordina el cobro y la entrega.'
)
r.font.size = Pt(9.5)
r.font.color.rgb = C_EMERALD
r.font.name = 'Consolas'
doc.add_paragraph()

add_page_break(doc)

# ─── 10.4 Reporte privado al admin ───────────────────
add_heading(doc, '10.4 Reporte privado al Admin (chat 1:1 con Juan)', level=2)
add_body(doc,
    'Ademas del grupo, el bot envia un reporte de KPIs al chat privado del administrador. '
    'Disenado para ejecutarse via Vercel Cron Job o por demanda con un comando /metricas.', size=10)
doc.add_paragraph()

add_bullet(doc, 'Funcion: sendAdminMetricsReport(snapshot) en src/lib/telegram/service.ts')
add_bullet(doc, 'Variable de entorno: TELEGRAM_ADMIN_ID = chat ID privado de Juan')
add_bullet(doc, 'Contenido del snapshot:')
add_bullet(doc, 'periodo (ej: "Hoy", "Ultimos 7 dias")', indent_level=1)
add_bullet(doc, 'facturado (suma ARS de pedidos cerrado_pagado)', indent_level=1)
add_bullet(doc, 'entregados (cantidad de pedidos cerrado_pagado)', indent_level=1)
add_bullet(doc, 'ticketPromedio (calculo)', indent_level=1)
add_bullet(doc, 'pendientes (pendiente + en_revision)', indent_level=1)
add_bullet(doc, 'enConflicto (status = en_conflicto)', indent_level=1)
doc.add_paragraph()

# Ejemplo de reporte
add_body(doc, 'EJEMPLO DE REPORTE PRIVADO:', bold=True, size=9.5)
t = doc.add_table(rows=1, cols=1)
t.alignment = WD_TABLE_ALIGNMENT.LEFT
set_cell_bg(t.rows[0].cells[0], RGBColor(0x1E, 0x10, 0x40))
p = t.rows[0].cells[0].paragraphs[0]
p.paragraph_format.left_indent = Cm(0.5)
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(6)
r = p.add_run(
    '📊 Metricas — Hoy\n'
    '\n'
    '💰 Facturado: $14.800.000\n'
    '✅ Entregados: 22\n'
    '📊 Ticket promedio: $672.727\n'
    '⏳ Pendientes: 5\n'
    '⚠️ En conflicto: 2'
)
r.font.size = Pt(9.5)
r.font.color.rgb = C_VIOLET
r.font.name = 'Consolas'
doc.add_paragraph()

# ─── 10.5 Variables de entorno ───────────────────────
add_heading(doc, '10.5 Variables de entorno (Vercel)', level=2)

env_vars = [
    ('TELEGRAM_BOT_TOKEN',       'Token del bot (creado con @BotFather)',                'CRITICA'),
    ('TELEGRAM_GROUP_ID',        'ID del grupo de ventas (negativo, ej: -1001234567890)', 'CRITICA'),
    ('TELEGRAM_ADMIN_ID',        'Chat ID privado de Juan (positivo)',                   'CRITICA'),
    ('SUPABASE_WEBHOOK_SECRET',  'Secreto compartido para validar webhooks',             'CRITICA'),
    ('SUPABASE_SERVICE_ROLE_KEY','Service-role key para queries internas',                'CRITICA'),
    ('NEXT_PUBLIC_SUPABASE_URL', 'URL del proyecto Supabase',                            'CRITICA'),
]

table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, lbl in enumerate(['VARIABLE', 'DESCRIPCION', 'NIVEL']):
    set_cell_bg(table.rows[0].cells[i], C_ZINC800)
    pp = table.rows[0].cells[i].paragraphs[0]
    pp.paragraph_format.space_before = Pt(4)
    pp.paragraph_format.space_after = Pt(4)
    rr = pp.add_run(lbl)
    rr.font.size = Pt(8.5)
    rr.font.bold = True
    rr.font.color.rgb = C_WHITE

for (name, desc, level) in env_vars:
    row = table.add_row()
    for j, (val, color, mono) in enumerate([
        (name, C_AMBER, True),
        (desc, C_ZINC400, False),
        (level, C_RED, False),
    ]):
        set_cell_bg(row.cells[j], C_DARK)
        pp = row.cells[j].paragraphs[0]
        pp.paragraph_format.space_before = Pt(4)
        pp.paragraph_format.space_after = Pt(4)
        pp.paragraph_format.left_indent = Cm(0.2)
        rr = pp.add_run(val)
        rr.font.size = Pt(8.5)
        rr.font.color.rgb = color
        if mono: rr.font.name = 'Consolas'

doc.add_paragraph()

# ─── 10.6 Troubleshooting ────────────────────────────
add_heading(doc, '10.6 Troubleshooting (debugging del webhook)', level=2)

add_body(doc, 'Si las notificaciones no llegan al grupo, seguir este checklist:', bold=True, size=10)
doc.add_paragraph()

add_bullet(doc, 'PASO 1 — Healthcheck publico:')
add_bullet(doc, 'GET https://tu-dominio.vercel.app/api/webhooks/supabase/telegram?ping=1', indent_level=1)
add_bullet(doc, 'Devuelve JSON con envChecks (cada env var) y bot.ok (si el token es valido).', indent_level=1)
add_bullet(doc, 'Si envChecks tiene algun false → falta esa env var en Vercel.', indent_level=1)
doc.add_paragraph()

add_bullet(doc, 'PASO 2 — Deployment Protection:')
add_bullet(doc, 'Vercel Dashboard → Settings → Deployment Protection.', indent_level=1)
add_bullet(doc, 'Si "Vercel Authentication" esta activa, Supabase recibe 401 HTML.', indent_level=1)
add_bullet(doc, 'Solucion: usar "Protection Bypass for Automation" y agregar el bypass token a la URL del webhook.', indent_level=1)
add_bullet(doc, 'Alternativa: agregar dominio custom (los dominios custom no estan protegidos).', indent_level=1)
doc.add_paragraph()

add_bullet(doc, 'PASO 3 — Test mode (sin pasar por DB):')
add_bullet(doc, 'POST con headers: x-supabase-signature: <secret>, x-test-mode: true', indent_level=1)
add_bullet(doc, 'Body: {"test":"ping"} → llama getMe a Telegram', indent_level=1)
add_bullet(doc, 'Body: {"test":"group-hello"} → manda un hola al grupo', indent_level=1)
add_bullet(doc, 'Body: {"test":"approved-mock"} → simula un evento "aprobado" sin tocar la DB', indent_level=1)
doc.add_paragraph()

add_bullet(doc, 'PASO 4 — Logs en Vercel:')
add_bullet(doc, 'Vercel Dashboard → Project → Logs', indent_level=1)
add_bullet(doc, 'Filtrar por prefijo [TG-WH] — cada paso del handler tiene su log.', indent_level=1)
add_bullet(doc, 'Buscar: "Telegram sendToGroup falló" o "lookup error" para identificar la falla exacta.', indent_level=1)
doc.add_paragraph()

add_bullet(doc, 'PASO 5 — Webhook configurado en Supabase:')
add_bullet(doc, 'Supabase Dashboard → Database → Webhooks → verificar que existan dos:', indent_level=1)
add_bullet(doc, '(a) order_events / INSERT  → URL del endpoint', indent_level=1)
add_bullet(doc, '(b) orders / UPDATE        → URL del endpoint', indent_level=1)
add_bullet(doc, 'Header: x-supabase-signature: valor EXACTO de SUPABASE_WEBHOOK_SECRET', indent_level=1)
doc.add_paragraph()

add_page_break(doc)

# ════════════════════════════════════════════════════
# SECCION 11 — RESUMEN EJECUTIVO
# ════════════════════════════════════════════════════
add_heading(doc, '11. RESUMEN EJECUTIVO', level=1)
add_body(doc, 'Cuadro sinoptico de los flujos por actor.', size=10)
doc.add_paragraph()

# Tabla resumen
table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, (lbl, col) in enumerate([('TALLER', C_EMERALD), ('VENDEDOR', C_ORANGE), ('ADMIN', C_VIOLET)]):
    set_cell_bg(table.rows[0].cells[i], C_ZINC800)
    p = table.rows[0].cells[i].paragraphs[0]
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(5)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(lbl)
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = col

acciones = [
    ([
        '+ Crear pedido',
        '+ Ver cotizacion',
        '+ Elegir metodo de pago',
        '+ Aprobar / Parcial / Rechazar',
        '+ Enviar WhatsApp al vendedor',
        '+ Iniciar reclamo (desde cerrado)',
        '+ Contactar por WA si hay conflicto',
        '+ Eliminar pedido (solo pendiente)',
    ], [
        '+ Tomar / Liberar pedidos',
        '+ Marcar en revision',
        '+ Armar y enviar cotizacion',
        '+ Marcar como Pagado',
        '+ Marcar como Entregado y Pagado',
        '+ Marcar como Entregado (desde pagado)',
        '+ Resolver conflicto (modal)',
        '+ Eliminar pedidos (pendiente/revision)',
    ], [
        '+ Todo lo del Vendedor',
        '+ Confirmar pago (cerrado → cerrado_pagado)',
        '+ Resolver conflictos',
        '+ Ver metricas globales',
        '+ Ver ranking vendedores/talleres',
        '+ Gestionar usuarios y roles',
        '+ Ver pedidos eliminados',
        '+ Reasignar vendedor',
    ]),
]

row = table.add_row()
for j, (acciones_col, color) in enumerate(zip(acciones[0], [C_EMERALD, C_ORANGE, C_VIOLET])):
    set_cell_bg(row.cells[j], C_DARK)
    for k, accion in enumerate(acciones_col):
        p = row.cells[j].paragraphs[0] if k == 0 else row.cells[j].add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        p.paragraph_format.left_indent  = Cm(0.3)
        r = p.add_run(accion)
        r.font.size = Pt(9)
        r.font.color.rgb = color

doc.add_paragraph()

add_separator(doc)

p_fin = doc.add_paragraph()
p_fin.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_fin = p_fin.add_run('Portal B2B Autopartes — Playbook v3.0 — Flujo + Telegram')
r_fin.font.size = Pt(8)
r_fin.font.color.rgb = C_ZINC600
r_fin.font.italic = True

# ─── Guardar ──────────────────────────────────────────────────────────────────
doc.save(OUTPUT_PATH)
print(f"OK Documento generado: {OUTPUT_PATH}")
